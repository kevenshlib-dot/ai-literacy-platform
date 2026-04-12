"""Parse exam papers from Word (.docx) format into the standard paper JSON format.

Handles typical Chinese exam paper structures:
- 判断题 (True/False)
- 单选题 / 多选题 (Single/Multiple Choice)
- 填空题 (Fill-in-the-blank)
- 简答题 / 论述题 / 开放题 (Short answer / Essay)

Extracts sections, questions, options, and answers (from answer key sections if present).
"""
import re
import logging
from typing import Optional
from docx import Document

logger = logging.getLogger(__name__)

# ── Section / question type detection ───────────────────────────────────────

# Multiple section header patterns (ordered by priority)
SECTION_PATTERNS = [
    # 第一部分：判断题（每题3分，共30分）
    re.compile(r"第[一二三四五六七八九十\d]+部分[：:]?\s*(.+)"),
    # 一、判断题（每题3分，共30分）  /  一．判断题  /  一. 判断题
    re.compile(r"^[一二三四五六七八九十]+[、．.]\s*(.+)"),
    # (一)、判断题   or  （一）判断题
    re.compile(r"^[（(][一二三四五六七八九十]+[）)][、．.]?\s*(.+)"),
    # I、判断题  II、单选题  (Roman-ish)
    re.compile(r"^[IⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩiⅰⅱⅲⅳⅴ]+[、．.]\s*(.+)"),
    # Part 1: ... / Part I: ...
    re.compile(r"^Part\s+[\dIVXivx]+[：:.]\s*(.+)", re.IGNORECASE),
]

# Primary keywords — short, unambiguous, matched with simple "in" check
QUESTION_TYPE_MAP = {
    "判断题": "true_false",
    "判断": "true_false",
    "是非题": "true_false",
    "是非": "true_false",
    "单选题": "single_choice",
    "单选": "single_choice",
    "多选题": "multiple_choice",
    "多选": "multiple_choice",
    "选择题": "single_choice",
    "填空题": "fill_blank",
    "填空": "fill_blank",
    "简答题": "short_answer",
    "简答": "short_answer",
    "论述题": "short_answer",
    "论述": "short_answer",
    "开放题": "short_answer",
    "主观题": "short_answer",
    "问答题": "short_answer",
    "问答": "short_answer",
    # Ambiguous keywords — only match with 题 suffix to avoid false positives
    "综合题": "short_answer",
    "分析题": "short_answer",
    "计算题": "short_answer",
    "案例题": "short_answer",
    "应用题": "short_answer",
    "True/False": "true_false",
    "Multiple Choice": "multiple_choice",
    "Single Choice": "single_choice",
}

QUESTION_NUM_PATTERN = re.compile(
    r"^(\d+)[.、．)\s]+(.+)",
)

OPTION_PATTERN = re.compile(
    r"^([A-Fa-f])[.、．)\s]+(.+)",
)

# Patterns for answer sections
ANSWER_SECTION_PATTERNS = [
    re.compile(r"^参考答案"),
    re.compile(r"^答案与"),
    re.compile(r"^标准答案"),
    re.compile(r"^评分标准"),
    re.compile(r"^答案[：:\s]"),
    re.compile(r"^【答案】"),
    re.compile(r"^【.*答���.*】"),  # e.g., 【第一部分：���非题答案】
    # "第 1 题：✓ 正确" / "第1题：✗ 错误" — inline answer+explanation format
    re.compile(r"^第\s*1\s*题\s*[：:]\s*[√×✓✗对错]"),
    re.compile(r"^第\s*1\s*题\s*[：:]\s*(?:正确|错误)"),
    re.compile(r"^第\s*1\s*题\s*[：:]\s*[A-Fa-f]"),
]

TF_ANSWER_PATTERN = re.compile(
    r"(?:第\s*)?(\d+)(?:\s*题)?\s*[.、．]*[：:]?\s*([√×✓✗对错TFXY]|正确|错误)",
)

MC_ANSWER_PATTERN = re.compile(
    r"(\d+)[.、．\s]*[：:]?\s*([A-Fa-f]+)",
)

# Bulk answer line like "1. B　　 2. C　　 3. B"
BULK_MC_PATTERN = re.compile(
    r"(\d+)[.、．\s]*([A-Fa-f]+)",
)


def _detect_question_type(text: str) -> Optional[str]:
    """Detect question type from section title text."""
    for keyword, qtype in QUESTION_TYPE_MAP.items():
        if keyword in text:
            return qtype
    return None


def _match_section_header(text: str) -> Optional[str]:
    """Try all section patterns, return the captured section title or None."""
    for pattern in SECTION_PATTERNS:
        m = pattern.match(text)
        if m:
            return m.group(1).strip()
    return None


def _is_standalone_type_header(text: str) -> Optional[str]:
    """Check if a line is a standalone type header like '判断题' or '判断题（每题3分）'.

    Returns the question type string or None.
    Must be strict to avoid matching titles or question content.
    """
    # Must be short enough to be a header
    if len(text) > 40:
        return None
    # Must not start with a number (which is a question line)
    if re.match(r"^\d+", text):
        return None
    # Must not start with an option letter
    if re.match(r"^[A-Fa-f][.、．)\s]", text):
        return None

    # The text must contain "X题" pattern (e.g., 判断题, 单选题, 简答题)
    # to distinguish from titles like "AI素养综合测试" that happen to contain "综合"
    type_header_pattern = re.compile(
        r"(?:判断|是非|单选|多选|选择|填空|简答|论述|开放|主观|问答|综合|分析|计算|案例)题"
    )
    if not type_header_pattern.search(text):
        return None

    qtype = _detect_question_type(text)
    if not qtype:
        return None

    # The type keyword should appear near the start (within first 8 chars)
    for keyword in QUESTION_TYPE_MAP:
        kw_with_ti = keyword + "题"
        if kw_with_ti in text:
            idx = text.find(kw_with_ti)
            if idx <= 8:
                return qtype
    return None


def _parse_score_from_text(text: str) -> tuple[Optional[float], Optional[int], Optional[float]]:
    """Extract per-question score, question count, and total score from text.

    Returns (per_q, count, total).
    Handles many common Chinese exam formats:
        "每题3分，共30分"       → (3.0, None, 30.0)
        "每小题3分，共30分"     → (3.0, None, 30.0)
        "每道题6分，共30分"     → (6.0, None, 30.0)
        "每道6分"              → (6.0, None, None)
        "共30分"               → (None, None, 30.0)
        "2分/题，共10分"        → (2.0, None, 10.0)
        "10题，每题3分"         → (3.0, 10, None)
        "3分×10题=30分"        → (3.0, 10, 30.0)
    """
    per_q = None
    total = None
    count = None

    # 每题X分 / 每小题X分 / 每道题X分 / 每道X分
    m_per = re.search(r"每(?:小题|道题|道|题)\s*(\d+(?:\.\d+)?)\s*分", text)
    if m_per:
        per_q = float(m_per.group(1))

    # X分/题 / X分每题
    if per_q is None:
        m_per2 = re.search(r"(\d+(?:\.\d+)?)\s*分[/／每](?:小题|道题|道|题)", text)
        if m_per2:
            per_q = float(m_per2.group(1))

    # X分×N题 / X分*N题
    if per_q is None:
        m_mul = re.search(r"(\d+(?:\.\d+)?)\s*分\s*[×xX*]\s*(\d+)\s*题", text)
        if m_mul:
            per_q = float(m_mul.group(1))
            count = int(m_mul.group(2))

    # 共X分 / 满分X分 / 合计X分
    m_total = re.search(r"(?:共|满分|合计)\s*(\d+(?:\.\d+)?)\s*分", text)
    if m_total:
        total = float(m_total.group(1))

    # X题 (question count) — only if not already set
    if count is None:
        m_count = re.search(r"(\d+)\s*(?:道)?题", text)
        if m_count:
            count = int(m_count.group(1))

    # Fallback: compute per_q from total and count
    if per_q is None and total is not None and count is not None and count > 0:
        per_q = round(total / count, 1)

    return per_q, count, total


def _is_score_info_line(text: str) -> bool:
    """Check if a line is purely score/instruction info (not a question or option).

    e.g. "每题3分，共30分" / "共10题，满分40分" / "本部分共5道题"
    """
    # Must be short
    if len(text) > 60:
        return False
    # Must not start with a question number
    if re.match(r"^\d+[.、．)\s]", text):
        return False
    # Must not start with an option letter
    if re.match(r"^[A-Fa-f][.、．)\s]", text):
        return False
    # Must contain score-related keywords
    if re.search(r"(?:分|题)", text):
        per_q, count, total = _parse_score_from_text(text)
        return per_q is not None or total is not None
    return False


def _normalize_tf_answer(ans: str) -> str:
    """Normalize true/false answer to standard T/F format."""
    a = ans.strip()
    if a in ("√", "✓", "对", "T", "Y", "正确", "t", "y"):
        return "T"
    elif a in ("×", "✗", "错", "F", "X", "N", "错误", "f", "x", "n"):
        return "F"
    return a


def _extract_title(paragraphs: list[str]) -> tuple[str, str, int]:
    """Extract title and description from the beginning of the document.

    Returns (title, description, content_start_idx).
    content_start_idx: index of first paragraph that is NOT part of the title/description area.
    """
    if not paragraphs:
        return ("试卷", "", 0)

    title = ""
    description = ""
    content_start = 0

    # Gather candidate title lines (before any section header or question)
    header_lines = []
    for i, text in enumerate(paragraphs):
        # Stop if we hit a section header, standalone type header, or first question
        if _match_section_header(text):
            content_start = i
            break
        if _is_standalone_type_header(text):
            content_start = i
            break
        if QUESTION_NUM_PATTERN.match(text):
            content_start = i
            break
        # Skip purely instructional lines
        if text.startswith("说明") or text.startswith("注意"):
            continue
        header_lines.append((i, text))
        content_start = i + 1

    if not header_lines:
        # No header lines found — probably starts directly with questions
        return (paragraphs[0], "", 0)

    # Heuristic: find the best title among header lines
    # Prefer a line that contains exam-related keywords
    exam_keywords = ["测验", "试卷", "考试", "试题", "测试", "考核", "检测", "考查", "评估", "素养"]

    best_title_idx = 0
    for idx, (_, text) in enumerate(header_lines):
        for kw in exam_keywords:
            if kw in text:
                best_title_idx = idx
                break

    title = header_lines[best_title_idx][1]

    # Gather remaining header lines as description
    desc_parts = []
    for idx, (_, text) in enumerate(header_lines):
        if idx != best_title_idx:
            desc_parts.append(text)
    description = " ".join(desc_parts) if desc_parts else ""

    return (title, description, content_start)


def _extract_answers_from_tables(doc) -> dict[str, dict[int, str]]:
    """Extract answers from Word document tables.

    Many Chinese exam papers store answer keys in structured tables with columns:
    ['题号', '正确答案', '解析要点'] or similar.

    This function:
    1. Finds the "参考答案" header table to identify where answer tables start
    2. Parses subsequent answer tables, auto-detecting question type from answer format
    3. Returns answers in the same format as paragraph-based answer parsing

    Returns:
        dict: section_type -> {question_num: "answer" or "answer|explanation"}
    """
    answers: dict[str, dict[int, str]] = {}

    if not doc.tables:
        return answers

    # Find the answer section start: a table containing "参考答案" or "标准答案"
    answer_section_start = None
    for i, table in enumerate(doc.tables):
        if len(table.rows) == 1:
            cell_text = table.rows[0].cells[0].text.strip()
            if any(kw in cell_text for kw in ("参考答案", "标准答案", "答案与解析")):
                answer_section_start = i
                break

    if answer_section_start is None:
        # No answer section header found; try to detect answer tables directly
        # by looking for tables with '题号' + '正确答案' header
        for i, table in enumerate(doc.tables):
            if len(table.rows) >= 2:
                header_cells = [c.text.strip() for c in table.rows[0].cells]
                if "题号" in header_cells and any(
                    kw in " ".join(header_cells) for kw in ("正确答案", "答案")
                ):
                    if answer_section_start is None:
                        answer_section_start = i - 1  # assume previous is header

    if answer_section_start is None:
        return answers

    # Also try to determine section types from header tables before the answer section
    # e.g., Table 3: "第一部分：是非题" → first answer table = true_false
    section_type_order = []
    for i in range(answer_section_start):
        table = doc.tables[i]
        if len(table.rows) == 1:
            cell_text = table.rows[0].cells[0].text.strip()
            detected = _detect_question_type(cell_text)
            if detected:
                section_type_order.append(detected)

    # Parse answer tables (those after the answer section header with '题号' header)
    answer_table_idx = 0
    for i in range(answer_section_start + 1, len(doc.tables)):
        table = doc.tables[i]
        if len(table.rows) < 2:
            continue

        header_cells = [c.text.strip() for c in table.rows[0].cells]
        if "题号" not in header_cells:
            continue

        # Find column indices
        num_col = header_cells.index("题号")
        ans_col = None
        exp_col = None
        for ci, h in enumerate(header_cells):
            if "答案" in h and ans_col is None:
                ans_col = ci
            if "解析" in h and exp_col is None:
                exp_col = ci

        if ans_col is None:
            continue

        # Parse data rows
        table_answers: dict[int, str] = {}
        detected_type = None

        for row in table.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) <= max(num_col, ans_col):
                continue

            qnum_str = cells[num_col]
            if not qnum_str.isdigit():
                continue
            qnum = int(qnum_str)

            raw_answer = cells[ans_col]
            explanation = cells[exp_col] if exp_col is not None and len(cells) > exp_col else ""

            # Detect question type from answer format
            if raw_answer in ("√", "✓", "对", "T", "Y", "正确"):
                if detected_type is None:
                    detected_type = "true_false"
                answer = "T"  # T = 正确
            elif raw_answer in ("×", "✗", "错", "F", "X", "N", "错误"):
                if detected_type is None:
                    detected_type = "true_false"
                answer = "F"  # F = 错误
            elif "、" in raw_answer or "," in raw_answer:
                # Multiple choice: "A、C、D" or "A,C,D"
                if detected_type is None:
                    detected_type = "multiple_choice"
                # Extract just the letters
                answer = re.sub(r"[^A-Fa-f]", "", raw_answer).upper()
            else:
                # Single letter or short answer
                clean = raw_answer.strip().upper()
                if re.match(r"^[A-F]$", clean):
                    if detected_type is None:
                        detected_type = "single_choice"
                    answer = clean
                else:
                    answer = raw_answer

            entry = answer
            if explanation:
                entry = answer + "|" + explanation

            table_answers[qnum] = entry

        # Determine final section type
        if detected_type is None and answer_table_idx < len(section_type_order):
            detected_type = section_type_order[answer_table_idx]

        if detected_type and table_answers:
            if detected_type not in answers:
                answers[detected_type] = {}
            answers[detected_type].update(table_answers)
            logger.info(
                f"Extracted {len(table_answers)} answers from table {i} "
                f"(type={detected_type})"
            )

        answer_table_idx += 1

    return answers


def _extract_ordered_text(doc) -> list[str]:
    """Extract all text from the document in document order, including both
    paragraphs and single-row table cells (which are often used as section headers).

    This ensures section headers stored in tables are not missed by the parser.
    """
    from docx.oxml.ns import qn

    texts = []
    body = doc.element.body
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            # Regular paragraph
            text = child.text or ""
            # Also get text from runs
            for r in child.iter(qn("w:t")):
                pass  # child.text already handled by python-docx
            # Use python-docx paragraph text extraction
            from docx.text.paragraph import Paragraph
            para = Paragraph(child, doc)
            text = para.text.strip()
            if text:
                texts.append(text)
        elif tag == "tbl":
            # Table: include small tables (≤3 rows) which are typically
            # title blocks, section headers, or instruction lines.
            # Multi-row answer tables (with '题号' header) are handled by
            # _extract_answers_from_tables separately.
            from docx.table import Table
            table = Table(child, doc)
            num_rows = len(table.rows)
            if num_rows <= 3:
                # Small table — extract first cell text from each row
                for row in table.rows:
                    cell_text = row.cells[0].text.strip()
                    if cell_text:
                        texts.append(cell_text)
    return texts


def parse_word_paper(file_data: bytes, filename: str = "paper.docx") -> dict:
    """Parse a Word document into our standard paper JSON format.

    Returns:
        dict in the standard paper import format:
        {
            "format_version": "1.0",
            "paper": {
                "title": str,
                "description": str,
                "sections": [...],
                "unsectioned_questions": [...],
            }
        }
    """
    import io
    doc = Document(io.BytesIO(file_data))

    # Extract text in document order, including single-row table cells
    paragraphs = _extract_ordered_text(doc)

    if not paragraphs:
        raise ValueError("Word文档内容为空")

    # 1. Extract title and description
    title, description, content_start = _extract_title(paragraphs)

    # Fallback title from filename
    if not title or len(title) < 2:
        title = filename.rsplit(".", 1)[0]

    # 2. Find answer section boundary
    answer_start_idx = None
    for i, text in enumerate(paragraphs):
        for pat in ANSWER_SECTION_PATTERNS:
            if pat.match(text):
                answer_start_idx = i
                break
        if answer_start_idx is not None:
            break

    # 3. Parse answers from the answer section
    answers: dict[str, dict[int, str]] = {}  # section_type -> {qnum: answer}
    current_answer_section_type = None

    if answer_start_idx is not None:
        answer_paragraphs = paragraphs[answer_start_idx:]
        for text in answer_paragraphs:
            is_header_line = not re.match(r"^\d+", text) and len(text) < 80

            if is_header_line:
                sec_title = _match_section_header(text)
                if sec_title:
                    detected = _detect_question_type(sec_title)
                    if detected:
                        current_answer_section_type = detected
                else:
                    standalone = _is_standalone_type_header(text)
                    if standalone:
                        current_answer_section_type = standalone

            # Auto-infer section type from line format when no explicit sub-header found
            if current_answer_section_type is None:
                # Strict TF symbols only (avoid confusing single-choice "B" etc.)
                if re.search(r"(?:第\s*)?(\d+)(?:\s*题)?\s*[.、．]*[：:]?\s*([√×✓✗]|正确|错误)", text):
                    current_answer_section_type = "true_false"
                # Auto-infer choice type from "1. B" / "1. ACD" patterns
                elif re.search(r"^\d+[.、．\s]*[A-Fa-f]+\s*$", text.strip()):
                    bulk = BULK_MC_PATTERN.findall(text)
                    if bulk:
                        # Check if any answer has multiple letters → multiple_choice
                        has_multi = any(len(a) > 1 for _, a in bulk)
                        current_answer_section_type = "multiple_choice" if has_multi else "single_choice"

            # Parse answer lines — try all formats regardless of current_answer_section_type
            # so that answers are never silently lost
            parsed_this_line = False

            # Try T/F format
            tf_match = TF_ANSWER_PATTERN.search(text)
            if tf_match:
                inferred_type = current_answer_section_type or "true_false"
                if inferred_type not in answers:
                    answers[inferred_type] = {}
                qnum = int(tf_match.group(1))
                ans = _normalize_tf_answer(tf_match.group(2))
                answers[inferred_type][qnum] = ans
                explanation_start = tf_match.end()
                explanation = text[explanation_start:].strip().lstrip("　 ")
                if explanation:
                    answers[inferred_type][qnum] = ans + "|" + explanation
                parsed_this_line = True

            # Try choice format (1. B / 1. ACD / bulk "1. B  2. C  3. B")
            if not parsed_this_line:
                bulk_matches = BULK_MC_PATTERN.findall(text)
                if bulk_matches:
                    has_multi = any(len(a) > 1 for _, a in bulk_matches)
                    inferred_type = current_answer_section_type
                    if inferred_type is None:
                        inferred_type = "multiple_choice" if has_multi else "single_choice"
                    if inferred_type not in answers:
                        answers[inferred_type] = {}
                    for qnum_str, ans in bulk_matches:
                        answers[inferred_type][int(qnum_str)] = ans.upper()
                    parsed_this_line = True

            # Try explanation format: "第X题：..."
            if current_answer_section_type and current_answer_section_type in answers:
                exp_match = re.search(r"第\s*(\d+)\s*题[：:]\s*(.+)", text)
                if exp_match:
                    qnum = int(exp_match.group(1))
                    if qnum in answers[current_answer_section_type]:
                        existing = answers[current_answer_section_type][qnum]
                        if "|" not in existing:
                            answers[current_answer_section_type][qnum] = existing + "|" + exp_match.group(2)

    # 3b. Extract answers from Word document tables (many papers store answers in tables)
    table_answers = _extract_answers_from_tables(doc)
    if table_answers:
        logger.info(f"Found answers in tables: {', '.join(f'{k}={len(v)}' for k, v in table_answers.items())}")
        # Merge table answers into paragraph answers; table answers take priority
        # (they tend to be more structured and reliable)
        for section_type, qnum_map in table_answers.items():
            if section_type not in answers:
                answers[section_type] = {}
            for qnum, ans_entry in qnum_map.items():
                # Only override if paragraph parsing didn't find an answer for this question
                if qnum not in answers[section_type] or not answers[section_type][qnum]:
                    answers[section_type][qnum] = ans_entry
                else:
                    # If paragraph had answer without explanation but table has explanation, merge
                    existing = answers[section_type][qnum]
                    if "|" not in existing and "|" in ans_entry:
                        answers[section_type][qnum] = ans_entry

    # 4. Parse question sections (before answer section)
    question_paragraphs = paragraphs[:answer_start_idx] if answer_start_idx else paragraphs

    sections = []
    current_section = None
    current_qtype = None
    current_score_per = 5.0
    current_question = None
    section_question_num = 0
    global_order = 1
    unsectioned_questions = []

    start_idx = content_start

    def _start_new_section(sec_title: str, full_text: str):
        """Create a new section and update current_qtype / current_score_per."""
        nonlocal current_section, current_qtype, current_score_per, section_question_num

        qtype = _detect_question_type(sec_title)
        per_q, count, total = _parse_score_from_text(full_text)

        current_qtype = qtype or "short_answer"
        current_score_per = per_q or 5.0
        section_question_num = 0

        current_section = {
            "title": sec_title.split("（")[0].split("(")[0].strip(),
            "description": sec_title if ("（" in sec_title or "(" in sec_title) else None,
            "order_num": len(sections) + 1,
            "score_rule": {"score_per_question": current_score_per, "total": total, "count": count},
            "questions": [],
        }
        sections.append(current_section)

    def _absorb_score_info(text: str):
        """Absorb a score-info line into the current section's score_rule.

        Called when we detect a score-info line right after a section header.
        Updates current_score_per and current_section's score_rule.
        """
        nonlocal current_score_per
        if not current_section:
            return
        per_q, count, total = _parse_score_from_text(text)
        rule = current_section.get("score_rule") or {}

        if per_q and (rule.get("score_per_question") == 5.0 or rule.get("score_per_question") is None):
            current_score_per = per_q
            rule["score_per_question"] = per_q
        if total and not rule.get("total"):
            rule["total"] = total
        if count and not rule.get("count"):
            rule["count"] = count

        # Re-compute per_q from total/count if still missing
        if rule.get("score_per_question") == 5.0 and rule.get("total") and rule.get("count"):
            computed = round(rule["total"] / rule["count"], 1)
            current_score_per = computed
            rule["score_per_question"] = computed

        current_section["score_rule"] = rule

    def _flush_question():
        nonlocal current_question, global_order, section_question_num
        if not current_question:
            return

        # Auto-infer question type from content if current_qtype seems wrong
        effective_qtype = current_qtype or "short_answer"
        options = current_question.get("options", {})

        if effective_qtype == "true_false" and len(options) > 2:
            # Has more than 2 options → likely single/multiple choice, not T/F
            effective_qtype = "single_choice"
        elif effective_qtype != "true_false" and effective_qtype != "fill_blank":
            # If it has standard A/B/C/D options, it's a choice question
            if len(options) >= 2:
                if effective_qtype == "short_answer":
                    effective_qtype = "single_choice"

        # Look up answer — try exact type first, then fallback search across all types
        answer_dict = answers.get(effective_qtype, {})
        answer_entry = answer_dict.get(section_question_num, "")

        # Fallback: if no answer found, search across all answer sections
        # This handles cases where the answer section type detection differs from
        # the question section type detection (e.g., answer parsed as "single_choice"
        # but question inferred as "multiple_choice", or answer section had no type header)
        if not answer_entry:
            for atype, adict in answers.items():
                if section_question_num in adict:
                    answer_entry = adict[section_question_num]
                    logger.debug(
                        f"Answer for Q{section_question_num} found via fallback "
                        f"(expected type={effective_qtype}, found in type={atype})"
                    )
                    break

        if "|" in answer_entry:
            correct_answer, explanation = answer_entry.split("|", 1)
        else:
            correct_answer = answer_entry
            explanation = None

        # For true/false, set standard options and normalize answer to T/F
        if effective_qtype == "true_false":
            if not options:
                current_question["options"] = {"T": "正确", "F": "错误"}
            # Normalize legacy A/B or text values
            if correct_answer in ("正确", "A", "√", "✓", "对"):
                correct_answer = "T"
            elif correct_answer in ("错误", "B", "×", "✗", "错"):
                correct_answer = "F"

        # Clean up empty options
        final_options = current_question.get("options") or None
        if isinstance(final_options, dict) and not final_options:
            final_options = None

        question_data = {
            "order_num": global_order,
            "score": current_score_per,
            "question": {
                "question_type": effective_qtype,
                "stem": current_question["stem"],
                "options": final_options,
                "correct_answer": correct_answer,
                "explanation": explanation,
                "difficulty": 3,
                "dimension": None,
                "knowledge_tags": None,
                "bloom_level": None,
            },
        }

        if current_section is not None:
            current_section["questions"].append(question_data)
        else:
            unsectioned_questions.append(question_data)

        global_order += 1
        current_question = None

    for text in question_paragraphs[start_idx:]:
        # ── Priority 1: Structured section header (一、..., 第X部分：...)
        sec_title = _match_section_header(text)
        if sec_title:
            _flush_question()
            _start_new_section(sec_title, text)
            continue

        # ── Priority 2: Standalone type header (判断题, 单选题（每题3分）, ...)
        #    IMPORTANT: always flush pending question first, so we don't
        #    accidentally skip this header because current_question is set.
        standalone_type = _is_standalone_type_header(text)
        if standalone_type:
            _flush_question()
            _start_new_section(text, text)
            continue

        # ── Priority 3: Score info line (e.g. "每题3分，共30分" on its own line)
        #    Only absorb if we just created a section and haven't started a question yet
        if current_section is not None and current_question is None and _is_score_info_line(text):
            _absorb_score_info(text)
            continue

        # ── Priority 4: Skip instruction lines
        if text.startswith("请判断") or text.startswith("请选择") or text.startswith("请用") or text.startswith("说明"):
            continue

        # ── Priority 5: Question number
        q_match = QUESTION_NUM_PATTERN.match(text)
        if q_match:
            _flush_question()
            section_question_num = int(q_match.group(1))
            stem = q_match.group(2).strip()
            current_question = {"stem": stem, "options": {}}
            continue

        # ── Priority 6: Option line (A. xxx, B. xxx, ...)
        opt_match = OPTION_PATTERN.match(text)
        if opt_match and current_question is not None:
            opt_letter = opt_match.group(1).upper()
            opt_text = opt_match.group(2).strip()
            current_question["options"][opt_letter] = opt_text
            continue

        # ── Priority 7: Multi-line stem continuation
        if current_question is not None:
            current_question["stem"] += "\n" + text

    # Flush last question
    _flush_question()

    # 5. Post-process: fix per-question scores using section totals
    for sec in sections:
        rule = sec.get("score_rule") or {}
        sec_total = rule.get("total")
        sec_per_q = rule.get("score_per_question")
        sec_count = rule.get("count")
        actual_count = len(sec["questions"])

        if actual_count == 0:
            continue

        # If per-question score is the default (5.0) and we have total,
        # recalculate from total / actual_count
        if sec_per_q == 5.0 and sec_total and actual_count > 0:
            computed_per_q = round(sec_total / actual_count, 1)
            for q in sec["questions"]:
                q["score"] = computed_per_q
            sec["score_rule"]["score_per_question"] = computed_per_q

    # 6. If no sections were detected but we have unsectioned questions,
    #    auto-group by question type
    if not sections and unsectioned_questions:
        type_groups: dict[str, list] = {}
        for q in unsectioned_questions:
            qt = q["question"]["question_type"]
            type_groups.setdefault(qt, []).append(q)

        type_labels = {
            "true_false": "判断题",
            "single_choice": "单选题",
            "multiple_choice": "多选题",
            "fill_blank": "填空题",
            "short_answer": "简答题",
        }

        order = 1
        for qt, qs in type_groups.items():
            sections.append({
                "title": type_labels.get(qt, "其他题型"),
                "description": None,
                "order_num": order,
                "score_rule": None,
                "questions": qs,
            })
            order += 1
        unsectioned_questions = []

    # 7. Calculate total score
    total_score = sum(
        sum(q["score"] for q in sec["questions"])
        for sec in sections
    )
    total_score += sum(q["score"] for q in unsectioned_questions)

    total_questions = sum(len(s["questions"]) for s in sections) + len(unsectioned_questions)

    # 8. Build result
    result = {
        "format_version": "1.0",
        "paper": {
            "title": title,
            "description": description if description != title else None,
            "total_score": total_score,
            "time_limit_minutes": None,
            "tags": None,
            "sections": sections,
            "unsectioned_questions": unsectioned_questions,
        },
    }

    logger.info(
        f"Parsed Word paper: {title}, {len(sections)} sections, "
        f"{total_questions} questions, total {total_score} pts"
    )

    if total_questions == 0:
        logger.warning(
            f"Word paper parsed with 0 questions! "
            f"Total paragraphs: {len(paragraphs)}, "
            f"Question paragraphs scanned: {len(question_paragraphs) - start_idx}"
        )

    # 9. Check for missing answers and generate warnings
    warnings = []
    objective_types = {"single_choice", "multiple_choice", "true_false"}
    all_questions_list = []
    for sec in sections:
        all_questions_list.extend(sec.get("questions", []))
    all_questions_list.extend(unsectioned_questions)

    missing_answer_count = 0
    for q_item in all_questions_list:
        qq = q_item.get("question", {})
        qtype = qq.get("question_type", "")
        answer = qq.get("correct_answer", "")
        if qtype in objective_types and not answer.strip():
            missing_answer_count += 1
            warnings.append(
                f"第{q_item.get('order_num', '?')}题（{qtype}）缺少正确答案，请在导入后手动设置"
            )

    if missing_answer_count > 0:
        logger.warning(
            f"Word paper has {missing_answer_count}/{total_questions} objective questions without correct answers"
        )

    result["warnings"] = warnings
    return result
