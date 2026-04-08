"""Export papers to Word (.docx) format for traditional exam use.

Generates a well-formatted Word document with:
- Title page with paper metadata
- Sections with section headers
- Questions formatted by type (true/false, choice, fill-blank, short answer)
- Optional answer key appended at the end
"""
import io
import logging
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

logger = logging.getLogger(__name__)

QUESTION_TYPE_LABELS = {
    "true_false": "判断题",
    "single_choice": "单选题",
    "multiple_choice": "多选题",
    "fill_blank": "填空题",
    "short_answer": "简答题",
}

SECTION_NUM_MAP = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]


def _get_section_num(idx: int) -> str:
    """Convert 0-based index to Chinese section number."""
    if idx < len(SECTION_NUM_MAP):
        return SECTION_NUM_MAP[idx]
    return str(idx + 1)


def _setup_styles(doc: Document):
    """Set up document styles for consistent formatting."""
    # Normal style
    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.5


def _add_title(doc: Document, title: str, subtitle: Optional[str] = None):
    """Add the paper title section."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "黑体"

    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(subtitle)
        run2.font.size = Pt(12)
        run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _add_metadata(doc: Document, time_limit: Optional[int], total_score: float):
    """Add metadata line (time limit, total score)."""
    parts = []
    if time_limit:
        parts.append(f"考试时间：{time_limit} 分钟")
    parts.append(f"满分：{total_score:.0f} 分")

    if parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("    ".join(parts))
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Add a separator line
    doc.add_paragraph("_" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_section_header(doc: Document, idx: int, title: str, description: Optional[str], score_rule: Optional[dict]):
    """Add a section header."""
    section_num = _get_section_num(idx)

    # Build the header text
    header_text = f"第{section_num}部分：{title}"

    if score_rule:
        score_per = score_rule.get("score_per_question")
        total = score_rule.get("total")
        score_parts = []
        if score_per:
            score_parts.append(f"每题 {score_per:.0f} 分")
        if total:
            score_parts.append(f"共 {total:.0f} 分")
        if score_parts:
            header_text += f"（{'，'.join(score_parts)}）"

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(header_text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = "黑体"

    # Add instruction text based on question type
    if description:
        p_desc = doc.add_paragraph()
        run_desc = p_desc.add_run(description)
        run_desc.font.size = Pt(10)
        run_desc.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _add_tf_question(doc: Document, num: int, stem: str):
    """Add a true/false question."""
    # Remove trailing bracket placeholders
    stem = stem.rstrip("（　）()").strip()
    p = doc.add_paragraph()
    run = p.add_run(f"{num}. {stem}（　　）")
    run.font.size = Pt(11)


def _add_choice_question(doc: Document, num: int, stem: str, options: Optional[dict]):
    """Add a choice question with options."""
    p = doc.add_paragraph()
    run = p.add_run(f"{num}. {stem}")
    run.font.size = Pt(11)

    if options:
        sorted_opts = sorted(options.items(), key=lambda x: x[0])
        for letter, text in sorted_opts:
            p_opt = doc.add_paragraph()
            p_opt.paragraph_format.left_indent = Cm(1)
            p_opt.paragraph_format.space_before = Pt(1)
            p_opt.paragraph_format.space_after = Pt(1)
            run_opt = p_opt.add_run(f"{letter}. {text}")
            run_opt.font.size = Pt(11)


def _add_short_answer_question(doc: Document, num: int, stem: str):
    """Add a short answer question with answer space."""
    p = doc.add_paragraph()
    run = p.add_run(f"{num}. {stem}")
    run.font.size = Pt(11)

    # Add blank lines for answer space
    for _ in range(4):
        blank = doc.add_paragraph()
        blank.paragraph_format.space_before = Pt(2)
        blank.paragraph_format.space_after = Pt(2)
        run_line = blank.add_run("_" * 70)
        run_line.font.size = Pt(10)
        run_line.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


def _add_fill_blank_question(doc: Document, num: int, stem: str):
    """Add a fill-in-the-blank question."""
    p = doc.add_paragraph()
    run = p.add_run(f"{num}. {stem}")
    run.font.size = Pt(11)


def _add_answer_key(doc: Document, sections: list, all_questions_flat: list):
    """Add answer key section at the end."""
    doc.add_page_break()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run("参考答案与评分标准")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "黑体"

    p_note = doc.add_paragraph()
    p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_note = p_note.add_run("（仅供阅卷使用）")
    run_note.font.size = Pt(10)
    run_note.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    for sec_idx, section in enumerate(sections):
        section_num = _get_section_num(sec_idx)
        p_sec = doc.add_paragraph()
        p_sec.paragraph_format.space_before = Pt(12)
        run_sec = p_sec.add_run(f"第{section_num}部分：{section['title']} 答案")
        run_sec.bold = True
        run_sec.font.size = Pt(12)

        questions = section.get("questions", [])
        qtype = None
        if questions:
            qtype = questions[0].get("question", {}).get("question_type")

        if qtype in ("single_choice", "multiple_choice"):
            # Compact format for choice questions
            line_items = []
            for i, q in enumerate(questions, 1):
                qd = q.get("question", {})
                ans = qd.get("correct_answer", "")
                line_items.append(f"{i}. {ans}")
                if len(line_items) == 5:
                    p_ans = doc.add_paragraph()
                    run_ans = p_ans.add_run("　　".join(line_items))
                    run_ans.font.size = Pt(11)
                    line_items = []
            if line_items:
                p_ans = doc.add_paragraph()
                run_ans = p_ans.add_run("　　".join(line_items))
                run_ans.font.size = Pt(11)

        elif qtype == "true_false":
            for i, q in enumerate(questions, 1):
                qd = q.get("question", {})
                ans = qd.get("correct_answer", "")
                explanation = qd.get("explanation", "")

                # Convert A/B back to √/×
                if ans == "A":
                    display_ans = "√"
                elif ans == "B":
                    display_ans = "×"
                else:
                    display_ans = ans

                p_ans = doc.add_paragraph()
                text = f"{i}. {display_ans}"
                if explanation:
                    text += f"　{explanation}"
                run_ans = p_ans.add_run(text)
                run_ans.font.size = Pt(11)

        elif qtype == "short_answer":
            p_note2 = doc.add_paragraph()
            run_note2 = p_note2.add_run("评分要点：")
            run_note2.bold = True
            run_note2.font.size = Pt(11)

            for i, q in enumerate(questions, 1):
                qd = q.get("question", {})
                explanation = qd.get("explanation", "")
                answer = qd.get("correct_answer", "")

                p_ans = doc.add_paragraph()
                text = f"第{i}题："
                if answer:
                    text += answer
                if explanation:
                    if answer:
                        text += "；"
                    text += explanation
                if not answer and not explanation:
                    text += "（开放题，评分参考具体评分标准）"
                run_ans = p_ans.add_run(text)
                run_ans.font.size = Pt(11)


def export_paper_to_word(paper_detail: dict, include_answers: bool = True) -> bytes:
    """Export a paper detail dict to a Word document.

    Args:
        paper_detail: Paper detail from get_paper_detail() or export format.
                     Should contain title, description, sections with questions.
        include_answers: Whether to append an answer key at the end.

    Returns:
        bytes: The .docx file content.
    """
    doc = Document()
    _setup_styles(doc)

    # Handle both export format and detail format
    paper_data = paper_detail.get("paper", paper_detail)

    title = paper_data.get("title", "试卷")
    description = paper_data.get("description")
    total_score = paper_data.get("total_score", 0)
    time_limit = paper_data.get("time_limit_minutes")
    sections = paper_data.get("sections", [])
    unsectioned = paper_data.get("unsectioned_questions", [])

    # Title
    _add_title(doc, title, description)
    _add_metadata(doc, time_limit, total_score)

    # Process each section
    question_counter = 0
    for sec_idx, section in enumerate(sections):
        sec_title = section.get("title", f"部分 {sec_idx + 1}")
        sec_desc = section.get("description")
        score_rule = section.get("score_rule")
        questions = section.get("questions", [])

        if not questions:
            continue

        # Detect question type from first question
        first_q = questions[0].get("question", {})
        qtype = first_q.get("question_type", "short_answer")

        # Add instruction based on type
        instructions = {
            "true_false": '请判断以下陈述是否正确，在括号内填写"\u221A"或"\u00D7"。',
            "single_choice": "请选择最佳答案。",
            "multiple_choice": "请选择所有正确答案。",
            "fill_blank": "请在横线上填写正确答案。",
            "short_answer": "请简要作答。",
        }

        _add_section_header(doc, sec_idx, sec_title, instructions.get(qtype), score_rule)

        # Add questions
        for i, q_item in enumerate(questions, 1):
            qd = q_item.get("question", {})
            q_type = qd.get("question_type", "short_answer")
            stem = q_item.get("stem_override") or qd.get("stem", "")
            options = q_item.get("options_override") or qd.get("options")

            if q_type == "true_false":
                _add_tf_question(doc, i, stem)
            elif q_type in ("single_choice", "multiple_choice"):
                _add_choice_question(doc, i, stem, options)
            elif q_type == "fill_blank":
                _add_fill_blank_question(doc, i, stem)
            else:
                _add_short_answer_question(doc, i, stem)

            question_counter += 1

    # Unsectioned questions
    if unsectioned:
        doc.add_paragraph().paragraph_format.space_before = Pt(16)
        for i, q_item in enumerate(unsectioned, question_counter + 1):
            qd = q_item.get("question", {})
            q_type = qd.get("question_type", "short_answer")
            stem = q_item.get("stem_override") or qd.get("stem", "")
            options = q_item.get("options_override") or qd.get("options")

            if q_type in ("single_choice", "multiple_choice"):
                _add_choice_question(doc, i, stem, options)
            elif q_type == "true_false":
                _add_tf_question(doc, i, stem)
            elif q_type == "fill_blank":
                _add_fill_blank_question(doc, i, stem)
            else:
                _add_short_answer_question(doc, i, stem)

    # Answer key
    if include_answers:
        has_answers = any(
            q.get("question", {}).get("correct_answer")
            for sec in sections
            for q in sec.get("questions", [])
        )
        if has_answers:
            _add_answer_key(doc, sections, [])

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
