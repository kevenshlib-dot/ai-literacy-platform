import io
import uuid
import zipfile
import pytest
from httpx import AsyncClient, ASGITransport
from sqlalchemy import text
from sqlalchemy.ext.asyncio import create_async_engine, async_sessionmaker, AsyncSession

from app.main import app
from app.core.database import Base, get_db
from app.core.config import settings
from app.services.user_service import init_roles
from app.services.parsing_service import (
    parse_material, chunk_text,
    PDFParser, EPUBParser, WordParser, MarkdownParser, HTMLParser, CSVParser, JSONParser,
)
from app.services.parse_worker import parse_and_store
from app.models.material import Material, MaterialFormat, MaterialStatus


# ---- Unit Tests for Parsers ----


def build_test_epub_bytes() -> bytes:
    container_xml = """<?xml version="1.0" encoding="UTF-8"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
  <rootfiles>
    <rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>
  </rootfiles>
</container>
"""
    content_opf = """<?xml version="1.0" encoding="UTF-8"?>
<package xmlns="http://www.idpf.org/2007/opf" version="3.0">
  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/">
    <dc:title>EPUB测试</dc:title>
  </metadata>
  <manifest>
    <item id="chapter2" href="chapter2.xhtml" media-type="application/xhtml+xml"/>
    <item id="chapter1" href="chapter1.xhtml" media-type="application/xhtml+xml"/>
  </manifest>
  <spine>
    <itemref idref="chapter1"/>
    <itemref idref="chapter2"/>
  </spine>
</package>
"""
    chapter1 = """<?xml version="1.0" encoding="UTF-8"?>
<html xmlns="http://www.w3.org/1999/xhtml">
  <body><h1>第一章</h1><p>这是EPUB第一章内容。</p></body>
</html>
"""
    chapter2 = """<?xml version="1.0" encoding="UTF-8"?>
<html xmlns="http://www.w3.org/1999/xhtml">
  <body><h1>第二章</h1><p>这是EPUB第二章内容。</p></body>
</html>
"""

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as archive:
        archive.writestr("mimetype", "application/epub+zip")
        archive.writestr("META-INF/container.xml", container_xml)
        archive.writestr("OEBPS/content.opf", content_opf)
        archive.writestr("OEBPS/chapter1.xhtml", chapter1)
        archive.writestr("OEBPS/chapter2.xhtml", chapter2)
    return buf.getvalue()

def test_markdown_parser():
    text = parse_material(b"# Hello\n\nThis is **markdown**.", "test.md", "markdown")
    assert "Hello" in text
    assert "markdown" in text


def test_markdown_parser_strips_nul_bytes():
    text = parse_material(b"# A\x00I\n\nCon\x00tent", "test.md", "markdown")
    assert "\x00" not in text
    assert "AI" in text
    assert "Content" in text


def test_markdown_parser_handles_utf16():
    text = parse_material("AI基础知识".encode("utf-16"), "test.md", "markdown")
    assert "\x00" not in text
    assert "AI基础知识" in text


def test_epub_parser():
    text = parse_material(build_test_epub_bytes(), "book.epub", "epub")
    assert "第一章" in text
    assert "第二章" in text
    assert text.index("第一章") < text.index("第二章")


def test_epub_parser_with_ebooklib():
    pytest.importorskip("ebooklib")

    parser = EPUBParser()
    text = parser._parse_with_ebooklib(build_test_epub_bytes(), "book.epub")
    assert "第一章" in text
    assert "第二章" in text
    assert text.index("第一章") < text.index("第二章")


def test_invalid_epub_parser():
    with pytest.raises(Exception):
        parse_material(b"not-an-epub", "broken.epub", "epub")


def test_epub_parser_falls_back_to_native(monkeypatch):
    parser = EPUBParser()

    monkeypatch.setattr(
        parser,
        "_parse_with_ebooklib",
        lambda file_data, filename: (_ for _ in ()).throw(RuntimeError("boom")),
    )

    text = parser.parse(build_test_epub_bytes(), "book.epub")
    assert "第一章" in text
    assert "第二章" in text


def test_html_parser():
    html = b"<html><body><h1>Title</h1><p>Content here</p><script>var x=1;</script></body></html>"
    text = parse_material(html, "page.html", "html")
    assert "Title" in text
    assert "Content here" in text
    assert "var x" not in text  # Script content should be excluded


def test_csv_parser():
    csv_data = b"name,age\nAlice,30\nBob,25"
    text = parse_material(csv_data, "data.csv", "csv")
    assert "Alice" in text
    assert "Bob" in text


def test_json_parser():
    json_data = b'{"name": "test", "value": 42}'
    text = parse_material(json_data, "config.json", "json")
    assert "test" in text
    assert "42" in text


def test_image_parser_placeholder():
    text = parse_material(b"\x89PNG fake", "photo.png", "image")
    assert "图片素材" in text


def test_video_parser_placeholder():
    text = parse_material(b"fake video", "clip.mp4", "video")
    assert "视频素材" in text


def test_audio_parser_placeholder():
    text = parse_material(b"fake audio", "sound.mp3", "audio")
    assert "音频素材" in text


def test_unsupported_format():
    with pytest.raises(ValueError, match="不支持"):
        parse_material(b"data", "file.xyz", "xyz")


# ---- Chunking Tests ----

def test_chunk_text_basic():
    text = "A" * 1000
    chunks = chunk_text(text, chunk_size=500, overlap=50)
    assert len(chunks) >= 2
    assert len(chunks[0]) == 500


def test_chunk_text_empty():
    assert chunk_text("") == []
    assert chunk_text("   ") == []


def test_chunk_text_small():
    chunks = chunk_text("Short text", chunk_size=500, overlap=50)
    assert len(chunks) == 1
    assert chunks[0] == "Short text"


def test_chunk_text_strips_nul_bytes():
    chunks = chunk_text("AB\x00CD", chunk_size=10, overlap=1)
    assert chunks == ["ABCD"]


# ---- PDF Parser Test (with real PyPDF2) ----

def test_pdf_parser():
    """Create a minimal valid PDF and test parsing."""
    pytest.importorskip("pymupdf")
    from PyPDF2 import PdfWriter
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    buf = io.BytesIO()
    writer.write(buf)
    pdf_bytes = buf.getvalue()

    parser = PDFParser()
    text = parser.parse(pdf_bytes, "blank.pdf")
    # Blank page may produce empty text
    assert isinstance(text, str)


def test_pdf_parser_falls_back_to_pypdf2(monkeypatch):
    from PyPDF2 import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    buf = io.BytesIO()
    writer.write(buf)
    pdf_bytes = buf.getvalue()

    parser = PDFParser()
    monkeypatch.setattr(
        parser,
        "_parse_with_pymupdf",
        lambda file_data: (_ for _ in ()).throw(RuntimeError("boom")),
    )

    text = parser.parse(pdf_bytes, "blank.pdf")
    assert isinstance(text, str)


# ---- Word Parser Test ----

def test_word_parser():
    """Create a minimal valid docx and test parsing."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("第一段内容")
    doc.add_paragraph("第二段内容")
    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    parser = WordParser()
    text = parser.parse(docx_bytes, "test.docx")
    assert "第一段内容" in text
    assert "第二段内容" in text


# ---- Integration Tests: Parse API ----

async def override_get_db():
    engine = create_async_engine(settings.DATABASE_URL, echo=False)
    session_factory = async_sessionmaker(engine, class_=AsyncSession, expire_on_commit=False)
    async with session_factory() as session:
        try:
            yield session
            await session.commit()
        except Exception:
            await session.rollback()
            raise
    await engine.dispose()


app.dependency_overrides[get_db] = override_get_db


@pytest.fixture(autouse=True)
async def setup_db():
    engine = create_async_engine(settings.DATABASE_URL, echo=False)
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)
    async with engine.begin() as conn:
        await conn.execute(text("TRUNCATE TABLE users CASCADE"))
        await conn.execute(text("TRUNCATE TABLE materials CASCADE"))
        await conn.execute(text("TRUNCATE TABLE knowledge_units CASCADE"))
    session_factory = async_sessionmaker(engine, class_=AsyncSession, expire_on_commit=False)
    async with session_factory() as session:
        await init_roles(session)
        await session.commit()
    await engine.dispose()
    yield


def get_client():
    return AsyncClient(transport=ASGITransport(app=app), base_url="http://test")


async def register_user(client, role="organizer"):
    import uuid
    unique = uuid.uuid4().hex[:8]
    resp = await client.post("/api/v1/auth/register", json={
        "username": f"user_{unique}",
        "email": f"{unique}@test.com",
        "password": "password123",
        "role": role,
    })
    return resp.json()["access_token"]


@pytest.mark.asyncio
async def test_manual_parse_markdown():
    """Upload a markdown file, trigger manual parse, check knowledge units."""
    async with get_client() as client:
        token = await register_user(client, "organizer")

        # Upload markdown
        md_content = ("# AI基础认知\n\n" + "人工智能是计算机科学的一个分支。" * 50 + "\n\n"
                      "## 机器学习\n\n" + "机器学习是AI的核心技术之一。" * 50)
        upload = await client.post(
            "/api/v1/materials",
            files={"file": ("ai_basics.md", md_content.encode(), "text/markdown")},
            data={"title": "AI基础教材"},
            headers={"Authorization": f"Bearer {token}"},
        )
        assert upload.status_code == 201
        mid = upload.json()["id"]

        # Trigger manual parse
        parse_resp = await client.post(
            f"/api/v1/materials/{mid}/parse",
            headers={"Authorization": f"Bearer {token}"},
        )
        assert parse_resp.status_code == 200
        assert parse_resp.json()["parsed"] is True

        # Check knowledge units
        ku_resp = await client.get(
            f"/api/v1/materials/{mid}/knowledge-units",
            headers={"Authorization": f"Bearer {token}"},
        )
        assert ku_resp.status_code == 200
        data = ku_resp.json()
        assert data["total_units"] > 0
        assert data["status"] == "parsed"


@pytest.mark.asyncio
async def test_manual_parse_word():
    """Upload a Word document, trigger parse, verify knowledge units."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("AI伦理与治理概述")
    for i in range(10):
        doc.add_paragraph(f"第{i+1}条伦理准则：确保人工智能系统的公平性和透明性。" * 5)
    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    async with get_client() as client:
        token = await register_user(client, "organizer")
        upload = await client.post(
            "/api/v1/materials",
            files={"file": ("ethics.docx", docx_bytes,
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={"title": "AI伦理教材"},
            headers={"Authorization": f"Bearer {token}"},
        )
        assert upload.status_code == 201
        mid = upload.json()["id"]

        parse_resp = await client.post(
            f"/api/v1/materials/{mid}/parse",
            headers={"Authorization": f"Bearer {token}"},
        )
        assert parse_resp.status_code == 200
        assert parse_resp.json()["parsed"] is True

        ku_resp = await client.get(
            f"/api/v1/materials/{mid}/knowledge-units",
            headers={"Authorization": f"Bearer {token}"},
        )
        assert ku_resp.status_code == 200
        assert ku_resp.json()["total_units"] > 0


@pytest.mark.asyncio
async def test_manual_parse_epub():
    async with get_client() as client:
        token = await register_user(client, "organizer")
        upload = await client.post(
            "/api/v1/materials",
            files={"file": ("book.epub", build_test_epub_bytes(), "application/epub+zip")},
            data={"title": "EPUB教材"},
            headers={"Authorization": f"Bearer {token}"},
        )
        assert upload.status_code == 201
        mid = upload.json()["id"]

        parse_resp = await client.post(
            f"/api/v1/materials/{mid}/parse",
            headers={"Authorization": f"Bearer {token}"},
        )
        assert parse_resp.status_code == 200
        assert parse_resp.json()["parsed"] is True

        ku_resp = await client.get(
            f"/api/v1/materials/{mid}/knowledge-units",
            headers={"Authorization": f"Bearer {token}"},
        )
        assert ku_resp.status_code == 200
        data = ku_resp.json()
        assert data["total_units"] > 0
        assert data["status"] == "parsed"


@pytest.mark.asyncio
async def test_parse_and_store_marks_failed_after_flush_error(monkeypatch):
    material = Material(
        id=uuid.uuid4(),
        title="故障素材",
        format=MaterialFormat.MARKDOWN,
        file_path="bucket/material.md",
        status=MaterialStatus.UPLOADED,
        uploaded_by=uuid.uuid4(),
    )

    class FakeResult:
        def scalar_one_or_none(self):
            return material

    class FakeSession:
        def __init__(self):
            self.flush_calls = 0
            self.rollback_calls = 0
            self.added = []

        async def execute(self, _query):
            return FakeResult()

        async def flush(self):
            self.flush_calls += 1
            if self.flush_calls == 2:
                raise RuntimeError("invalid byte sequence for encoding UTF8: 0x00")

        async def rollback(self):
            self.rollback_calls += 1

        def add(self, obj):
            self.added.append(obj)

    async def fake_fetch_file_from_minio(_file_path: str) -> bytes:
        return b"broken"

    monkeypatch.setattr(
        "app.services.parse_worker.fetch_file_from_minio",
        fake_fetch_file_from_minio,
    )
    monkeypatch.setattr(
        "app.services.parse_worker.parse_material",
        lambda *_args, **_kwargs: "AI基础知识",
    )
    monkeypatch.setattr(
        "app.services.parse_worker.chunk_text",
        lambda *_args, **_kwargs: ["chunk-1"],
    )

    session = FakeSession()
    result = await parse_and_store(session, material.id)

    assert result is False
    assert session.rollback_calls == 1
    assert material.status == MaterialStatus.FAILED


@pytest.mark.asyncio
async def test_parse_material_not_found():
    async with get_client() as client:
        token = await register_user(client, "organizer")
        resp = await client.post(
            "/api/v1/materials/00000000-0000-0000-0000-000000000000/parse",
            headers={"Authorization": f"Bearer {token}"},
        )
    assert resp.status_code == 404


@pytest.mark.asyncio
async def test_auto_parse_on_upload():
    """Verify that uploading a CSV triggers background parsing automatically."""
    async with get_client() as client:
        token = await register_user(client, "organizer")
        upload = await client.post(
            "/api/v1/materials",
            files={"file": ("test.csv", b"a,b\n1,2", "text/csv")},
            data={"title": "自动解析素材"},
            headers={"Authorization": f"Bearer {token}"},
        )
        mid = upload.json()["id"]

        # Background parse runs after response; wait briefly
        import asyncio
        await asyncio.sleep(0.5)

        resp = await client.get(
            f"/api/v1/materials/{mid}/knowledge-units",
            headers={"Authorization": f"Bearer {token}"},
        )
    assert resp.status_code == 200
    # Auto-parsing should have created knowledge units
    assert resp.json()["total_units"] >= 1
