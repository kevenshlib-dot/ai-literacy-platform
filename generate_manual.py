"""生成 AI素养评测平台 完整系统说明书 (Word .docx)"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ── helpers ──────────────────────────────────────────────────────────
def set_cell_shading(cell, color_hex: str):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shd)

def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '1F4E79')
    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'E8F0FE')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def h1(doc, text):
    doc.add_heading(text, level=1)

def h2(doc, text):
    doc.add_heading(text, level=2)

def h3(doc, text):
    doc.add_heading(text, level=3)

def p(doc, text, bold=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(10.5)
    return para

def bullet(doc, text, level=0):
    para = doc.add_paragraph(text, style='List Bullet')
    para.paragraph_format.left_indent = Cm(1.27 + level * 1.27)
    for run in para.runs:
        run.font.size = Pt(10)
    return para

def note_box(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(f"📌 {text}")
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x8B, 0x45, 0x13)
    run.italic = True
    return para


# ── 开始构建文档 ─────────────────────────────────────────────────────
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 设置页边距
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)


# ═══════════════════════════════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run("AI素养评测平台")
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("系统说明书")
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x2A, 0x6B, 0xA6)

doc.add_paragraph()

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run("基于多智能体协同架构的智能化评测系统")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

for _ in range(6):
    doc.add_paragraph()

info_para = doc.add_paragraph()
info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info_para.add_run("版本：V1.0\n文档状态：正式发布")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
# 目录页（占位）
# ═══════════════════════════════════════════════════════════════════
toc_title = doc.add_paragraph()
toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = toc_title.add_run("目  录")
run.bold = True
run.font.size = Pt(18)

p(doc, "")
p(doc, "（请在 Word 中使用「引用 → 目录」功能自动生成目录）")
doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
# 第一部分：系统概述
# ═══════════════════════════════════════════════════════════════════
h1(doc, "第一部分  系统概述")

# ── 1. 平台简介 ──
h2(doc, "1. 平台简介")
p(doc, "AI素养评测平台是一套基于多智能体协同架构的智能化评测系统，旨在为教育机构、企业及个人提供全方位的人工智能素养能力评估与提升服务。")
p(doc, "平台核心理念：")
bullet(doc, "评测智能化：利用大语言模型（LLM）实现智能出题、智能评分、智能反馈")
bullet(doc, "学习个性化：基于自适应学习算法，为每位用户定制专属学习路径")
bullet(doc, "评估多维化：涵盖AI基础知识、技术应用、伦理安全、批判思维、创新实践五大维度")
bullet(doc, "流程闭环化：形成「评测 → 诊断 → 训练 → 再评测」的闭环提升体系")

p(doc, "平台名称：AI素养评测平台", bold=True)
p(doc, "版本号：V0.1.0")
p(doc, "默认访问地址：前端 http://localhost:3000 / 后端 API http://localhost:8000")

# ── 2. 系统架构 ──
h2(doc, "2. 系统架构概览")
p(doc, "系统采用前后端分离的微服务架构，核心由以下层次组成：")
bullet(doc, "展示层（前端）：Vue 3 + Ant Design Vue，提供响应式用户界面")
bullet(doc, "接口层（API Gateway）：FastAPI RESTful API，JWT认证，RBAC权限控制")
bullet(doc, "业务逻辑层（服务层）：30+ 业务服务，处理核心业务逻辑")
bullet(doc, "智能体层（Agent Layer）：8 个AI智能体，负责智能出题、评分、交互等")
bullet(doc, "数据层：PostgreSQL（关系数据）+ Milvus（向量数据）+ MinIO（文件存储）+ Elasticsearch（全文搜索）")
bullet(doc, "基础设施层：Redis（缓存）+ RabbitMQ（消息队列）+ Docker（容器化）")

p(doc, "")
p(doc, "系统数据流：", bold=True)
p(doc, "用户浏览器 → Nginx(80/443) → FastAPI(8000) → 业务服务 → 数据库/AI服务")
p(doc, "                                              ↕")
p(doc, "                                    LLM服务(DeepSeek/本地模型)")

# ── 3. 技术栈 ──
h2(doc, "3. 技术栈")
add_table(doc,
    ["层次", "技术", "版本", "用途"],
    [
        ["前端框架", "Vue 3", "3.5", "用户界面构建"],
        ["UI组件库", "Ant Design Vue", "4.2", "界面组件"],
        ["状态管理", "Pinia", "3.0", "前端状态管理"],
        ["构建工具", "Vite", "7.3", "前端构建与开发服务"],
        ["图表库", "ECharts", "6.0", "数据可视化"],
        ["后端框架", "FastAPI", "0.115", "异步 REST API"],
        ["ORM", "SQLAlchemy", "2.0", "数据库对象映射"],
        ["数据库", "PostgreSQL", "16", "关系型数据存储"],
        ["向量数据库", "Milvus", "2.3", "语义向量搜索"],
        ["对象存储", "MinIO", "-", "文件/素材存储"],
        ["搜索引擎", "Elasticsearch", "8.12", "全文搜索"],
        ["缓存", "Redis", "7", "缓存与任务队列"],
        ["消息队列", "RabbitMQ", "3.13", "异步消息处理"],
        ["AI/LLM", "OpenAI兼容API", "-", "智能出题、评分、交互"],
        ["认证", "JWT + bcrypt", "-", "用户认证与密码加密"],
        ["容器化", "Docker Compose", "-", "服务编排与部署"],
    ],
    col_widths=[3, 4, 2, 6]
)

# ── 4. 角色与权限 ──
h2(doc, "4. 角色与权限体系")
p(doc, "系统设有四种用户角色，每种角色拥有不同的功能权限：")

add_table(doc,
    ["角色", "中文名", "说明"],
    [
        ["admin", "管理员", "系统最高权限，管理用户、素材、全部功能"],
        ["organizer", "组织者", "创建考试、管理题库、查看成绩（注册需审批）"],
        ["reviewer", "审题员", "审核题目、评分答卷、查看成绩（注册需审批）"],
        ["examinee", "被测者", "参加考试、查看个人成绩、练习（注册即激活）"],
    ],
    col_widths=[3, 3, 9]
)

p(doc, "功能权限矩阵：", bold=True)
add_table(doc,
    ["功能模块", "管理员", "组织者", "审题员", "被测者"],
    [
        ["首页/工作台", "✓", "✓", "✓", "✗"],
        ["素材管理", "✓", "✗", "✗", "✗"],
        ["题库管理", "✓", "✓", "✓", "✗"],
        ["考试管理", "✓", "✓", "✗", "✗"],
        ["在线考试", "✓", "✓", "✓", "✓"],
        ["成绩管理（全部）", "✓", "✓", "✓", "✗"],
        ["成绩管理（个人）", "✓", "✓", "✓", "✓"],
        ["用户管理", "✓", "✗", "✗", "✗"],
        ["创建题目", "✓", "✓", "✗", "✗"],
        ["审核题目", "✓", "✗", "✓", "✗"],
        ["AI智能出题", "✓", "✓", "✗", "✗"],
        ["诊断报告", "✓", "✓", "✓", "✓"],
        ["证书下载", "✓", "✓", "✓", "✓"],
    ],
    col_widths=[4, 2, 2, 2, 2]
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
# 第二部分：用户操作指南
# ═══════════════════════════════════════════════════════════════════
h1(doc, "第二部分  用户操作指南")

# ── 5. 登录与注册 ──
h2(doc, "5. 登录与注册")

h3(doc, "5.1 登录")
p(doc, "访问平台首页，系统将显示登录页面。")
bullet(doc, "输入用户名和密码")
bullet(doc, "点击「登录」按钮")
bullet(doc, "登录成功后，管理员/组织者/审题员跳转至工作台，被测者跳转至在线考试页面")
note_box(doc, "默认管理员账号：admin / admin123（首次使用后请立即修改密码）")

h3(doc, "5.2 注册")
p(doc, "新用户可通过注册页面创建账号。")
bullet(doc, "填写用户名、邮箱、密码（至少6位）、确认密码")
bullet(doc, "选择注册角色：被测者、组织者、审题员")
bullet(doc, "被测者：注册后立即激活，可直接使用")
bullet(doc, "组织者/审题员：注册后需等待管理员审批通过后方可使用")
note_box(doc, "不允许通过注册页面注册管理员角色，管理员账号由系统初始化或已有管理员手动创建。")

# ── 6. 首页与工作台 ──
h2(doc, "6. 首页与工作台")

h3(doc, "6.1 首页")
p(doc, "首页展示平台整体运营概况，适用于管理员、组织者和审题员。")
p(doc, "统计卡片（顶部）：", bold=True)
bullet(doc, "素材总量：已上传的学习素材数量")
bullet(doc, "题库题目：题库中的题目总数")
bullet(doc, "考试场次：已创建的考试总数")
bullet(doc, "参与人数：平台注册用户总数")
p(doc, "点击任意卡片可快速跳转到对应的管理页面。")

p(doc, "快捷操作：", bold=True)
bullet(doc, "上传素材 → 跳转素材管理")
bullet(doc, "管理题库 → 跳转题库管理")
bullet(doc, "创建考试 → 跳转考试管理")

p(doc, "英雄榜：", bold=True)
p(doc, "首页底部展示成绩排行榜前5名，显示用户名、得分率和能力等级，并可点击「查看完整榜单」查看前20名详情。")

h3(doc, "6.2 工作台")
p(doc, "工作台为管理视角的仪表板，提供更详细的管理信息：")
bullet(doc, "统计卡片：素材、题目、考试、用户数量")
bullet(doc, "AI素养维度覆盖率：五大维度的进度条展示")
bullet(doc, "近期考试：最新5场考试的状态和参与情况")
bullet(doc, "待审核题目：最新5道等待审批的题目")
bullet(doc, "最新素材：最新5个上传的素材信息")

# ── 7. 素材管理 ──
h2(doc, "7. 素材管理")
note_box(doc, "本模块仅管理员可访问。")

h3(doc, "7.1 素材列表")
p(doc, "素材管理页面展示所有已上传的学习素材，支持筛选和搜索。")
p(doc, "筛选条件：", bold=True)
bullet(doc, "标题关键词搜索")
bullet(doc, "状态筛选：已上传、解析中、已解析、已向量化、失败")
bullet(doc, "格式筛选：PDF、Word、Markdown、图片、视频、音频、CSV、JSON")
bullet(doc, "分类筛选")

p(doc, "列表显示字段：", bold=True)
add_table(doc,
    ["字段", "说明"],
    [
        ["标题", "素材名称，可点击查看详情"],
        ["格式", "文件格式标签（彩色标识）"],
        ["文件大小", "自动转换为 B / KB / MB 显示"],
        ["状态", "处理状态标签"],
        ["分类", "素材所属分类"],
        ["上传时间", "格式：YYYY-MM-DD HH:mm"],
        ["操作", "详情、下载、删除"],
    ],
    col_widths=[3, 12]
)

h3(doc, "7.2 上传素材")
p(doc, "点击「上传素材」按钮打开上传窗口：")
bullet(doc, "标题（必填）：素材名称")
bullet(doc, "描述（选填）：素材描述信息")
bullet(doc, "分类（选填）：素材分类标签")
bullet(doc, "标签（选填）：逗号分隔的多个标签")
bullet(doc, "文件（必填）：支持拖拽上传")

p(doc, "支持的文件格式及大小限制：", bold=True)
add_table(doc,
    ["格式", "扩展名", "最大文件大小"],
    [
        ["PDF", ".pdf", "100 MB"],
        ["Word", ".doc, .docx", "50 MB"],
        ["Markdown", ".md, .markdown", "10 MB"],
        ["HTML", ".html, .htm", "10 MB"],
        ["图片", ".jpg, .jpeg, .png, .gif, .bmp, .webp", "20 MB"],
        ["视频", ".mp4, .avi, .mov, .wmv, .flv, .mkv, .webm", "500 MB"],
        ["音频", ".mp3, .wav, .flac, .aac, .ogg, .wma, .m4a", "200 MB"],
        ["CSV", ".csv", "50 MB"],
        ["JSON", ".json", "50 MB"],
    ],
    col_widths=[3, 6, 3]
)

h3(doc, "7.3 批量上传")
p(doc, "点击「批量上传」按钮可同时上传多个文件：")
bullet(doc, "可选择统一分类")
bullet(doc, "支持拖拽多文件上传")
bullet(doc, "上传完成后显示成功/失败数量及错误详情")

h3(doc, "7.4 素材处理流程")
p(doc, "素材上传后经历以下处理阶段：")
bullet(doc, "已上传 → 系统接收文件，存储至MinIO对象存储")
bullet(doc, "解析中 → 后台异步解析文件内容，提取知识单元")
bullet(doc, "已解析 → 知识单元提取完成，可在详情中查看")
bullet(doc, "已向量化 → 知识单元已生成向量嵌入，支持语义搜索")
note_box(doc, "素材详情面板（右侧抽屉）可查看素材的完整元数据和提取的知识单元列表。")

h3(doc, "7.5 语义搜索")
p(doc, "素材向量化后，支持通过自然语言进行语义相似度搜索，帮助快速定位相关知识内容。")

# ── 8. 题库管理 ──
h2(doc, "8. 题库管理")
p(doc, "题库管理是平台的核心功能之一，支持手动创建、AI智能生成、批量操作和审核流程。")

h3(doc, "8.1 题目类型")
add_table(doc,
    ["题型", "英文标识", "答题方式"],
    [
        ["单选题", "single_choice", "从A/B/C/D中选择一个正确答案"],
        ["多选题", "multiple_choice", "从A/B/C/D中选择多个正确答案"],
        ["判断题", "true_false", "判断对（A）或错（B）"],
        ["填空题", "fill_blank", "在文本框中输入答案"],
        ["简答题", "short_answer", "在文本框中输入详细回答"],
    ],
    col_widths=[3, 4, 8]
)

h3(doc, "8.2 题目筛选")
p(doc, "题库列表支持多维度筛选：")
bullet(doc, "关键词搜索（按题干内容）")
bullet(doc, "状态：草稿、待审核、已通过、已拒绝、已归档")
bullet(doc, "题型：单选/多选/判断/填空/简答")
bullet(doc, "难度：1-入门、2-简单、3-中等、4-困难、5-专家")
bullet(doc, "知识维度：按AI素养维度筛选")

h3(doc, "8.3 手动创建题目")
p(doc, "点击「新建题目」按钮，填写以下信息：")
bullet(doc, "题型（必选）：从下拉菜单选择")
bullet(doc, "题干（必填）：题目内容")
bullet(doc, "选项（选择题必填）：A、B、C、D选项内容")
bullet(doc, "正确答案（必填）：标准答案")
bullet(doc, "解析（选填）：答案解释说明")
bullet(doc, "难度（必选）：1-5星评级")
bullet(doc, "知识维度（选填）：所属AI素养维度")
bullet(doc, "布鲁姆认知层次（选填）：记忆、理解、应用、分析、评价、创造")
bullet(doc, "标签（选填）：知识点标签")

h3(doc, "8.4 AI智能出题")
p(doc, "平台支持三种AI出题方式：", bold=True)

p(doc, "方式一：基于知识单元生成", bold=True)
bullet(doc, "选择素材的某个知识单元")
bullet(doc, "指定题型、数量（1-10）、难度、认知层次")
bullet(doc, "AI自动生成题目")

p(doc, "方式二：基于素材批量生成", bold=True)
bullet(doc, "选择一个已解析的素材")
bullet(doc, "设定题型分布（如：单选10题、判断5题）")
bullet(doc, "AI遍历素材的知识单元批量生成")

p(doc, "方式三：自由生成（无素材）", bold=True)
bullet(doc, "不依赖已有素材")
bullet(doc, "直接利用LLM的知识库生成")
bullet(doc, "支持自定义提示词引导生成方向")

note_box(doc, "AI生成的题目默认为「草稿」状态，需提交审核并通过后方可用于组卷。当LLM服务不可用时，系统自动降级为模板生成。")

h3(doc, "8.5 审核流程")
p(doc, "题目从创建到可用遵循以下审核流程：")
bullet(doc, "草稿 → 提交审核 → 待审核 → 通过/拒绝")
bullet(doc, "管理员和审题员可执行审核操作")
bullet(doc, "支持单题审核和批量审核")
bullet(doc, "可通过AI质量检查辅助审核决策")

h3(doc, "8.6 批量操作")
p(doc, "选中多道题目后，可执行：")
bullet(doc, "批量提交审核")
bullet(doc, "批量通过")
bullet(doc, "批量拒绝")
bullet(doc, "批量删除")
bullet(doc, "导出题库（Markdown格式）")

h3(doc, "8.7 题库导入导出")
p(doc, "支持 Markdown 格式的题库导入和导出，便于题目的迁移和共享。")

# ── 9. 考试管理 ──
h2(doc, "9. 考试管理")
note_box(doc, "本模块仅管理员和组织者可访问。")

h3(doc, "9.1 创建考试")
p(doc, "点击「创建考试」按钮，填写：")
bullet(doc, "考试名称（必填）")
bullet(doc, "考试描述（选填）")
bullet(doc, "时间限制（选填）：1-300分钟，不填则不限时")
bullet(doc, "总分（必填）：1-1000分")

h3(doc, "9.2 智能组卷")
p(doc, "创建考试后，可使用智能组卷功能自动填充试题：")
bullet(doc, "设定题型分布：分别指定各题型的数量")
bullet(doc, "目标难度：通过滑杆选择1-5的难度目标")
bullet(doc, "难度容差：允许偏离目标难度的范围（0-2）")
bullet(doc, "每题分值：统一设定每题的分数")
bullet(doc, "知识维度：可选择特定维度的题目")
p(doc, "系统将从已通过审核的题库中，按照设定条件自动筛选和组合试题。")
note_box(doc, "也可通过「手动组卷」逐题添加。还支持自然语言描述需求，AI自动解析参数并组卷。")

h3(doc, "9.3 考试状态管理")
add_table(doc,
    ["状态", "说明", "可执行操作"],
    [
        ["草稿", "创建后的初始状态，可编辑", "编辑、组卷、发布、删除"],
        ["已发布", "考生可参加考试", "关闭"],
        ["已关闭", "考试结束，归档保存", "重新发布、删除"],
    ],
    col_widths=[3, 5, 6]
)

h3(doc, "9.4 考试分析")
p(doc, "对已有答题数据的考试，可查看分析报告：")
bullet(doc, "经典测量理论（CTT）分析：题目难度、区分度")
bullet(doc, "信度分析：Cronbach's α 系数")
bullet(doc, "题型分布统计")

# ── 10. 在线考试 ──
h2(doc, "10. 在线考试")
p(doc, "所有角色的用户均可参加在线考试。")

h3(doc, "10.1 考试列表")
p(doc, "在线考试页面展示所有已发布的考试，每个考试显示：")
bullet(doc, "考试名称和描述")
bullet(doc, "时间限制（如有）")
bullet(doc, "总分")
bullet(doc, "「开始考试」按钮（点击后需确认）")

h3(doc, "10.2 随机测试")
p(doc, "点击「测试一下？」可发起快速随机测试：")
bullet(doc, "设定题目数量：5-50题（步进5）")
bullet(doc, "选择难度模式：")
bullet(doc, "  自信心爆棚 —— 全部简单题", level=1)
bullet(doc, "  真实水平 —— 均衡难度，测试真实能力", level=1)
bullet(doc, "  挑战地狱难度 —— 高难度题目", level=1)
p(doc, "随机测试按 单选:多选:判断 = 6:2:2 比例抽题，满分100分，不限时。")

h3(doc, "10.3 答题界面")
p(doc, "进入考试后，界面分为两栏：")

p(doc, "左侧 - 题目导航：", bold=True)
bullet(doc, "网格式导航按钮，按题号排列")
bullet(doc, "白色 = 未答、浅蓝 = 已答、黄色边框 = 已标记、深蓝 = 当前题")

p(doc, "右侧 - 答题区域：", bold=True)
bullet(doc, "题号、题型标签、分值标签")
bullet(doc, "题干内容")
bullet(doc, "答题区（根据题型不同）：单选-单选按钮、多选-复选框、判断-对错按钮、填空/简答-文本框")
bullet(doc, "标记/取消标记按钮")
bullet(doc, "上一题/下一题导航")

p(doc, "顶部信息栏：", bold=True)
bullet(doc, "考试名称")
bullet(doc, "倒计时（剩余不足5分钟时红色闪烁提醒）")
bullet(doc, "答题进度（X/Y 已答）")
bullet(doc, "提交按钮（需二次确认）")

note_box(doc, "答案实时自动保存至服务器，切换题目不会丢失已填答案。考试超时将自动提交。")

h3(doc, "10.4 随机测试结果")
p(doc, "随机测试提交后立即显示结果：")
bullet(doc, "得分（X / Y）")
bullet(doc, "能力等级标签")
bullet(doc, "答对题目数和正确率")
bullet(doc, "返回考试列表按钮")

# ── 11. 成绩管理 ──
h2(doc, "11. 成绩管理")

h3(doc, "11.1 排行榜")
p(doc, "成绩页面顶部展示英雄榜：")
bullet(doc, "前三名显示金银铜牌标识")
bullet(doc, "展示用户名、得分率、等级、考试名称")
bullet(doc, "用户可通过「参与排名」开关选择是否出现在排行榜中")
bullet(doc, "完整榜单显示前20名")

h3(doc, "11.2 管理员/组织者视角")
p(doc, "管理员和组织者可查看所有成绩：")
bullet(doc, "搜索：按用户名、姓名、考试名称搜索")
bullet(doc, "成绩表字段：用户名、角色、考试名称、得分、得分率、等级、状态、提交时间")
bullet(doc, "操作：诊断报告、下载证书（优秀等级）、删除")
bullet(doc, "成绩导出：支持导出为 Excel (.xlsx) 文件")
bullet(doc, "归档管理：已删除的成绩可在归档视图中恢复")

h3(doc, "11.3 被测者/审题员视角")
p(doc, "查看个人成绩列表：")
bullet(doc, "考试名称、得分、等级、百分位排名、状态、提交时间")
bullet(doc, "操作：诊断分析报告、复盘（重新查看答题）、手动触发评分、删除")

h3(doc, "11.4 诊断分析报告")
p(doc, "点击「诊断报告」进入详细的评测分析页面：")

p(doc, "概览卡片：", bold=True)
bullet(doc, "总得分、得分率、能力等级、答对数量、百分位排名、用时")

p(doc, "五维度分析：", bold=True)
bullet(doc, "AI基础知识、AI技术应用、AI伦理安全、AI批判思维、AI创新实践")
bullet(doc, "柱状图展示各维度得分情况")

p(doc, "逐题分析表：", bold=True)
bullet(doc, "每道题的题干、用户答案、正确答案、得分、对错状态")

p(doc, "改进建议：", bold=True)
bullet(doc, "优势领域（高分维度）")
bullet(doc, "薄弱环节（低分维度）")
bullet(doc, "个性化学习建议")

h3(doc, "11.5 证书下载")
p(doc, "达到「优秀」等级的用户可下载电子证书，证书包含：姓名、日期、等级、得分、平台名称。")

h3(doc, "11.6 评分等级标准")
add_table(doc,
    ["等级", "得分率范围", "颜色标识"],
    [
        ["优秀", "≥ 90%", "金色"],
        ["良好", "75% - 89%", "蓝色"],
        ["合格", "60% - 74%", "绿色"],
        ["不合格", "< 60%", "红色"],
    ],
    col_widths=[3, 5, 4]
)

# ── 12. 用户管理 ──
h2(doc, "12. 用户管理")
note_box(doc, "本模块仅管理员可访问。")

h3(doc, "12.1 用户列表")
p(doc, "用户管理页面展示所有注册用户，支持筛选：")
bullet(doc, "关键词搜索（用户名/邮箱/姓名）")
bullet(doc, "角色筛选：管理员/组织者/被测者/审题员")
bullet(doc, "状态筛选：已启用/待审批/禁用")

p(doc, "列表字段及操作：", bold=True)
add_table(doc,
    ["字段/操作", "说明"],
    [
        ["用户名", "账号名称"],
        ["姓名", "用户真实姓名"],
        ["邮箱", "注册邮箱"],
        ["角色", "彩色标签显示（红=管理员，蓝=组织者，绿=被测者，橙=审题员）"],
        ["状态", "已启用（绿）/ 待审批（橙）/ 禁用（红）"],
        ["创建时间", "账号创建时间"],
        ["通过审批", "仅待审批用户可见，点击激活账号"],
        ["编辑", "修改用户姓名、电话、组织等信息"],
        ["重置密码", "管理员为用户设置新密码"],
        ["启用/禁用", "切换账号激活状态（不可操作自己）"],
        ["删除", "软删除用户（归档，可恢复）"],
    ],
    col_widths=[4, 11]
)

h3(doc, "12.2 添加用户")
p(doc, "管理员可手动添加用户，需填写：用户名、邮箱、密码、姓名（选填）、电话（选填）、组织（选填）、角色。")

h3(doc, "12.3 批量导入")
p(doc, "支持通过 Excel (.xlsx) 或 CSV 文件批量导入用户：")
bullet(doc, "必填列：username（用户名）、email（邮箱）")
bullet(doc, "导入用户统一分配「被测者」角色")
bullet(doc, "初始密码统一为 abcdefg")
bullet(doc, "可下载导入模板文件")
bullet(doc, "导入完成后显示成功/失败数量及错误详情")

h3(doc, "12.4 归档与恢复")
p(doc, "已删除的用户进入归档视图，管理员可：")
bullet(doc, "恢复用户：重新激活账号")
bullet(doc, "查看成绩：查看该用户的历史考试记录")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
# 第三部分：高级功能
# ═══════════════════════════════════════════════════════════════════
h1(doc, "第三部分  高级功能")

# ── 13. 交互式场景问答 ──
h2(doc, "13. 交互式场景问答")
p(doc, "交互式问答模块提供基于真实场景的多轮对话式评测体验。")
bullet(doc, "场景创建：系统生成特定AI应用场景")
bullet(doc, "多轮对话：用户与AI进行多轮交互，展示AI素养能力")
bullet(doc, "意图识别：AI实时识别用户意图并提供针对性反馈")
bullet(doc, "过程评分：不仅评价最终答案，还评估思考过程和交互质量")
bullet(doc, "难度自适应：根据用户表现动态调整问题难度")

p(doc, "API接口：", bold=True)
add_table(doc,
    ["操作", "说明"],
    [
        ["创建会话", "POST /interactive - 开始新的场景问答"],
        ["提交回答", "POST /interactive/{id}/respond - 提交本轮回答"],
        ["过程评分", "GET /interactive/{id}/process-score - 获取过程评分"],
        ["结束会话", "POST /interactive/{id}/end - 结束问答会话"],
    ],
    col_widths=[4, 11]
)

# ── 14. 沙盒练习 ──
h2(doc, "14. 沙盒练习环境")
p(doc, "沙盒提供安全的练习空间，不计入正式成绩。")
bullet(doc, "任务列表：浏览可用的练习任务")
bullet(doc, "创建会话：选择任务开始练习")
bullet(doc, "多次尝试：每次尝试都会获得AI反馈")
bullet(doc, "即时反馈：每次提交后获得针对性建议")
bullet(doc, "完成总结：会话结束时获得综合评价")
bullet(doc, "练习统计：查看个人练习数据和进步趋势")

# ── 15. 自适应学习 ──
h2(doc, "15. 自适应学习路径")
p(doc, "系统根据用户的评测表现，自动生成个性化学习路径。")

h3(doc, "15.1 弱项分析")
p(doc, "系统分析用户在五大维度上的表现，识别薄弱环节和知识盲区。")

h3(doc, "15.2 路径生成")
p(doc, "基于弱项分析结果，自动生成多步骤的学习路径：")
bullet(doc, "每个步骤包含学习目标和推荐资源")
bullet(doc, "步骤按照知识递进关系排序")
bullet(doc, "可追踪完成进度")

h3(doc, "15.3 课程推荐")
p(doc, "根据用户能力画像，智能推荐相关课程和学习材料。")

# ── 16. 闭环评估 ──
h2(doc, "16. 闭环评估-训练系统")
p(doc, "闭环系统实现「评测 → 诊断 → 训练 → 再评测」的完整学习循环：")
bullet(doc, "学习旅程：可视化展示用户的完整评测和训练历程")
bullet(doc, "前后对比：对比不同时期的评测结果，量化进步幅度")
bullet(doc, "效果跟踪：追踪训练对各维度能力提升的实际效果")
bullet(doc, "平台统计：管理员可查看全平台的闭环运营数据")

# ── 17. 指标框架 ──
h2(doc, "17. 指标框架管理")
p(doc, "指标框架定义了AI素养的评估维度和标准，支持动态更新。")
bullet(doc, "三智能体流水线：研究代理 → 分析代理 → 综合代理，协同生成指标提案")
bullet(doc, "提案审批：管理员审核AI生成的指标更新提案")
bullet(doc, "维度管理：管理五大AI素养维度及其子指标")

# ── 18. 组织管理 ──
h2(doc, "18. 组织管理")
p(doc, "支持多租户的组织架构管理：")
bullet(doc, "创建组织：管理员创建新的组织单位")
bullet(doc, "成员管理：添加/移除组织成员")
bullet(doc, "组织配置：自定义组织级别的参数设置")
bullet(doc, "组织统计：查看组织内的评测数据概览")

# ── 19. 课程管理 ──
h2(doc, "19. 课程管理")
p(doc, "管理员和组织者可创建结构化的学习课程：")
bullet(doc, "课程创建：设定课程名称、描述、维度、难度")
bullet(doc, "章节管理：在课程下添加章节，支持文本和视频内容")
bullet(doc, "发布管理：课程经历 草稿 → 已发布 → 已归档 的生命周期")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
# 第四部分：技术文档
# ═══════════════════════════════════════════════════════════════════
h1(doc, "第四部分  技术文档")

# ── 20. 系统架构 ──
h2(doc, "20. 系统架构详解")

h3(doc, "20.1 分层架构")
add_table(doc,
    ["层次", "职责", "关键技术"],
    [
        ["展示层", "用户界面、交互体验", "Vue 3, Ant Design Vue, ECharts"],
        ["接口层", "HTTP API、认证鉴权、路由", "FastAPI, JWT, CORS"],
        ["服务层", "业务逻辑、数据处理", "30+ Python 服务模块"],
        ["智能体层", "AI能力封装", "8个LLM Agent"],
        ["数据访问层", "ORM、数据持久化", "SQLAlchemy 2.0, asyncpg"],
        ["数据层", "数据存储", "PostgreSQL, Milvus, MinIO, ES"],
        ["基础设施层", "缓存、消息、容器", "Redis, RabbitMQ, Docker"],
    ],
    col_widths=[3, 5, 7]
)

h3(doc, "20.2 目录结构")
p(doc, "后端项目结构：")
bullet(doc, "app/main.py — 应用入口，中间件注册")
bullet(doc, "app/core/ — 核心配置（config.py, security.py, database.py）")
bullet(doc, "app/api/v1/endpoints/ — 18个API端点模块")
bullet(doc, "app/api/deps.py — 依赖注入（认证、角色检查）")
bullet(doc, "app/models/ — 15个SQLAlchemy数据模型")
bullet(doc, "app/schemas/ — Pydantic输入输出模型")
bullet(doc, "app/services/ — 30+业务服务")
bullet(doc, "app/agents/ — 8个AI智能体")
bullet(doc, "alembic/ — 数据库迁移")
bullet(doc, "tests/ — 测试用例")

p(doc, "前端项目结构：")
bullet(doc, "frontend/src/views/ — 10个页面组件")
bullet(doc, "frontend/src/router/ — 路由配置与权限守卫")
bullet(doc, "frontend/src/stores/ — Pinia状态管理")
bullet(doc, "frontend/src/layouts/ — 布局组件")
bullet(doc, "frontend/src/utils/ — 工具函数（API请求封装）")
bullet(doc, "frontend/src/components/ — 可复用组件")

# ── 21. API接口文档 ──
h2(doc, "21. API接口文档")
p(doc, "系统共有18个API模块，以下列出所有接口端点。基础路径：/api/v1")

h3(doc, "21.1 认证模块 (/auth)")
add_table(doc,
    ["方法", "路径", "权限", "说明"],
    [
        ["POST", "/auth/register", "公开", "用户注册"],
        ["POST", "/auth/login", "公开", "用户登录，返回JWT令牌"],
    ],
    col_widths=[2, 4, 2, 6]
)

h3(doc, "21.2 用户管理 (/users)")
add_table(doc,
    ["方法", "路径", "权限", "说明"],
    [
        ["GET", "/users/me", "已认证", "获取当前用户信息"],
        ["POST", "/users/me/reset-password", "已认证", "用户自助修改密码"],
        ["GET", "/users", "管理员", "用户列表（支持筛选）"],
        ["POST", "/users", "管理员", "手动创建用户"],
        ["POST", "/users/batch-import", "管理员", "批量导入用户"],
        ["PUT", "/users/{user_id}", "管理员", "更新用户信息"],
        ["POST", "/users/{user_id}/reset-password", "管理员", "重置用户密码"],
        ["DELETE", "/users/{user_id}", "管理员", "软删除用户"],
        ["POST", "/users/{user_id}/restore", "管理员", "恢复已删除用户"],
        ["GET", "/users/{user_id}/scores", "管理员", "查看用户成绩"],
    ],
    col_widths=[2, 5, 2, 5]
)

h3(doc, "21.3 素材管理 (/materials)")
add_table(doc,
    ["方法", "路径", "权限", "说明"],
    [
        ["POST", "/materials", "管理员/组织者", "上传素材"],
        ["POST", "/materials/batch", "管理员/组织者", "批量上传"],
        ["GET", "/materials", "已认证", "素材列表（筛选）"],
        ["GET", "/materials/{id}", "已认证", "素材详情"],
        ["GET", "/materials/{id}/download", "已认证", "获取下载链接"],
        ["POST", "/materials/{id}/parse", "管理员/组织者", "手动触发解析"],
        ["GET", "/materials/{id}/knowledge-units", "已认证", "知识单元列表"],
        ["POST", "/materials/{id}/vectorize", "管理员/组织者", "向量化知识单元"],
        ["GET", "/materials/search/semantic", "已认证", "语义搜索"],
        ["GET", "/materials/coverage", "管理员/组织者", "维度覆盖分析"],
        ["DELETE", "/materials/{id}", "管理员/组织者", "删除素材"],
    ],
    col_widths=[2, 5, 3, 4]
)

h3(doc, "21.4 题库管理 (/questions)")
add_table(doc,
    ["方法", "路径", "权限", "说明"],
    [
        ["POST", "/questions", "管理员/组织者", "创建题目"],
        ["GET", "/questions", "已认证", "题目列表"],
        ["GET", "/questions/stats", "已认证", "题库统计"],
        ["GET", "/questions/{id}", "已认证", "题目详情"],
        ["PUT", "/questions/{id}", "管理员/组织者", "更新题目"],
        ["DELETE", "/questions/{id}", "管理员", "删除题目"],
        ["POST", "/questions/{id}/submit", "管理员/组织者", "提交审核"],
        ["POST", "/questions/{id}/review", "管理员/审题员", "审核题目"],
        ["POST", "/questions/{id}/ai-check", "管理员/审题员", "AI质量检查"],
        ["POST", "/questions/generate", "管理员/组织者", "从知识单元生成"],
        ["POST", "/questions/generate/material/{id}", "管理员/组织者", "从素材批量生成"],
        ["POST", "/questions/generate/free", "管理员/组织者", "自由生成（无素材）"],
        ["POST", "/questions/generate/bank/{id}", "管理员/组织者", "按题型分布组建"],
        ["POST", "/questions/batch/submit", "管理员/组织者", "批量提交审核"],
        ["POST", "/questions/batch/review", "管理员/审题员", "批量审核"],
        ["POST", "/questions/batch/delete", "管理员", "批量删除"],
        ["POST", "/questions/batch/export-md", "管理员/组织者", "导出Markdown"],
        ["POST", "/questions/batch/import-md", "管理员/组织者", "导入Markdown"],
        ["POST", "/questions/calibration/auto-flag", "管理员", "自动标记低质量题"],
        ["POST", "/questions/calibration/{id}", "管理员/组织者", "重新校准难度"],
        ["GET", "/questions/calibration/similar", "管理员/组织者", "查找相似题（去重）"],
        ["GET", "/questions/analysis/report", "管理员/组织者", "全局质量报告"],
    ],
    col_widths=[2, 5, 3, 4]
)

h3(doc, "21.5 考试管理 (/exams)")
add_table(doc,
    ["方法", "路径", "权限", "说明"],
    [
        ["POST", "/exams", "管理员/组织者", "创建考试"],
        ["GET", "/exams", "已认证", "考试列表"],
        ["GET", "/exams/{id}", "已认证", "考试详情"],
        ["PUT", "/exams/{id}", "管理员/组织者", "更新考试"],
        ["DELETE", "/exams/{id}", "管理员", "删除考试"],
        ["POST", "/exams/{id}/publish", "管理员/组织者", "发布考试"],
        ["POST", "/exams/{id}/close", "管理员/组织者", "关闭考试"],
        ["POST", "/exams/{id}/reactivate", "管理员", "重新发布"],
        ["POST", "/exams/{id}/assemble/auto", "管理员/组织者", "智能组卷"],
        ["POST", "/exams/{id}/assemble/manual", "管理员/组织者", "手动组卷"],
        ["POST", "/exams/intent/assemble", "管理员/组织者", "自然语言组卷"],
        ["GET", "/exams/{id}/analysis", "管理员/组织者", "考试分析"],
    ],
    col_widths=[2, 5, 3, 4]
)

h3(doc, "21.6 考试会话 (/sessions)")
add_table(doc,
    ["方法", "路径", "权限", "说明"],
    [
        ["POST", "/sessions/start/{exam_id}", "已认证", "开始考试"],
        ["POST", "/sessions/random-test", "已认证", "创建随机测试"],
        ["GET", "/sessions/{sheet_id}", "已认证", "获取答题卡"],
        ["POST", "/sessions/{sheet_id}/answer", "已认证", "提交/更新答案"],
        ["POST", "/sessions/{sheet_id}/mark", "已认证", "标记题目"],
        ["POST", "/sessions/{sheet_id}/submit", "已认证", "提交考试"],
        ["GET", "/sessions", "已认证", "会话列表"],
    ],
    col_widths=[2, 5, 2, 5]
)

h3(doc, "21.7 成绩管理 (/scores)")
add_table(doc,
    ["方法", "路径", "权限", "说明"],
    [
        ["GET", "/scores/leaderboard", "已认证", "排行榜"],
        ["POST", "/scores/leaderboard/opt-out", "已认证", "切换排行榜可见"],
        ["GET", "/scores/all", "管理员/组织者", "全部成绩列表"],
        ["POST", "/scores/export", "管理员/组织者", "导出成绩Excel"],
        ["POST", "/scores/grade/{sheet_id}", "管理员/组织者/审题员", "评分答卷"],
        ["POST", "/scores/panel-score", "管理员/组织者/审题员", "多模型面板评分"],
        ["GET", "/scores/{score_id}", "已认证", "成绩详情"],
        ["GET", "/scores/{score_id}/diagnostic", "已认证", "五维诊断报告"],
        ["POST", "/scores/{score_id}/report", "已认证", "生成详细报告"],
        ["GET", "/scores/{score_id}/evaluation", "已认证", "评价反馈与排名"],
        ["POST", "/scores/training/generate", "已认证", "生成训练题目"],
        ["DELETE", "/scores/{sheet_id}", "已认证", "删除成绩"],
        ["GET", "/scores/archived", "管理员", "归档成绩列表"],
        ["POST", "/scores/{sheet_id}/restore", "管理员", "恢复成绩"],
    ],
    col_widths=[2, 5, 3, 4]
)

h3(doc, "21.8 其他模块")
add_table(doc,
    ["模块", "路径前缀", "主要功能"],
    [
        ["交互式问答", "/interactive", "场景创建、多轮对话、过程评分"],
        ["沙盒练习", "/sandbox", "练习任务、会话管理、尝试提交"],
        ["自适应学习", "/learning", "弱项分析、路径生成、推荐"],
        ["闭环评估", "/closed-loop", "学习旅程、前后对比、平台统计"],
        ["指标管理", "/indicators", "指标提案生成、审批"],
        ["标注管理", "/annotations", "自动/手动标注、一致性检查"],
        ["报告管理", "/reports", "月度运营报告"],
        ["组织管理", "/organizations", "组织CRUD、成员管理"],
        ["课程管理", "/courses", "课程与章节管理"],
        ["健康检查", "/health", "服务健康状态"],
    ],
    col_widths=[3, 4, 8]
)

# ── 22. 数据模型 ──
h2(doc, "22. 数据库模型")
p(doc, "系统核心数据表：")
add_table(doc,
    ["表名", "说明", "主要字段"],
    [
        ["users", "用户表", "id, username, email, hashed_password, role_id, is_active, is_deleted"],
        ["roles", "角色表", "id, name (admin/organizer/reviewer/examinee)"],
        ["materials", "素材表", "id, title, format, file_path, file_size, status, category, tags"],
        ["knowledge_units", "知识单元表", "id, material_id, title, content, chunk_index, embedding"],
        ["questions", "题目表", "id, question_type, stem, options, correct_answer, difficulty, status"],
        ["exams", "考试表", "id, title, description, total_score, time_limit, status"],
        ["exam_questions", "考试题目关联表", "id, exam_id, question_id, sequence, points"],
        ["answer_sheets", "答题卡表", "id, user_id, exam_id, status, started_at, submitted_at"],
        ["answers", "答案表", "id, sheet_id, question_id, user_answer, is_correct, score"],
        ["scores", "成绩表", "id, sheet_id, total_score, score_ratio, level, dimension_scores"],
        ["score_details", "成绩明细表", "id, score_id, question_id, 各维度得分"],
        ["reports", "报告表", "id, report_type, content, generated_at"],
    ],
    col_widths=[4, 4, 7]
)
note_box(doc, "所有表使用UUID主键，支持软删除（is_deleted + deleted_at），时间字段带时区。")

# ── 23. AI智能体 ──
h2(doc, "23. AI智能体体系")
p(doc, "系统内置8个专业AI智能体，各司其职：")

add_table(doc,
    ["智能体", "文件", "核心功能"],
    [
        ["评分智能体", "scoring_agent.py", "多模型评委面板评分（2-5个模型），位置交换去偏，维度化打分（准确性/完整性/逻辑/表达）"],
        ["出题智能体", "question_agent.py", "基于知识单元和布鲁姆分类法生成多类型题目"],
        ["素材智能体", "material_agent.py", "智能解析文档、提取知识单元、生成摘要"],
        ["交互智能体", "interactive_agent.py", "管理场景对话、分析回答、提供反馈"],
        ["意图智能体", "intent_agent.py", "解析自然语言描述，提取考试组卷参数"],
        ["标注智能体", "annotation_agent.py", "自动标注素材内容，维度和难度标记"],
        ["审核智能体", "review_agent.py", "AI质量审查题目，提供改进建议"],
        ["指标智能体", "indicator_agents.py", "三智能体流水线，动态更新评估维度指标"],
    ],
    col_widths=[3, 4, 8]
)

p(doc, "LLM配置说明：", bold=True)
bullet(doc, "云端模式：使用DeepSeek API（推荐），需设置API密钥")
bullet(doc, "本地模式：使用LM Studio或Ollama运行本地模型（需GPU）")
bullet(doc, "降级机制：LLM不可用时，自动切换为规则/模板模式")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
# 第五部分：部署与运维
# ═══════════════════════════════════════════════════════════════════
h1(doc, "第五部分  部署与运维")

# ── 24. 环境要求 ──
h2(doc, "24. 环境要求")
add_table(doc,
    ["项目", "最低配置", "推荐配置"],
    [
        ["操作系统", "Ubuntu 20.04 / CentOS 8 / macOS", "Ubuntu 22.04 LTS"],
        ["CPU", "4核", "8核"],
        ["内存", "8 GB", "16 GB"],
        ["硬盘", "50 GB SSD", "100 GB SSD"],
        ["Docker", "24.0+", "最新版"],
        ["Docker Compose", "2.0+", "最新版"],
        ["Node.js", "18+", "20 LTS"],
        ["Python", "3.10+", "3.11"],
        ["GPU（本地LLM）", "16GB VRAM", "24GB VRAM"],
    ],
    col_widths=[4, 5, 5]
)

# ── 25. 快速启动 ──
h2(doc, "25. 快速启动指南")
p(doc, "开发环境4步快速启动：", bold=True)

p(doc, "第1步：启动基础设施服务", bold=True)
bullet(doc, "cd ai-literacy-platform")
bullet(doc, "docker compose up -d")
bullet(doc, "等待所有容器启动完成（约1-2分钟）")

p(doc, "第2步：启动LLM服务（可选）", bold=True)
bullet(doc, "方式A：LM Studio → 启动本地模型，监听端口1234")
bullet(doc, "方式B：Ollama → ollama pull qwen2.5:14b && ollama serve")
bullet(doc, "方式C：云端API → 在.env中配置DeepSeek API密钥")
note_box(doc, "LLM服务非必须，不启动则AI出题/评分功能降级为模板模式。")

p(doc, "第3步：启动后端", bold=True)
bullet(doc, "pip install -r requirements.txt")
bullet(doc, "alembic upgrade head（首次需要）")
bullet(doc, "uvicorn app.main:app --reload --host 0.0.0.0 --port 8000")

p(doc, "第4步：启动前端", bold=True)
bullet(doc, "cd frontend")
bullet(doc, "npm install（首次需要）")
bullet(doc, "npm run dev")

p(doc, "启动后访问 http://localhost:3000 即可使用平台。", bold=True)

# ── 26. Docker部署 ──
h2(doc, "26. Docker服务说明")
add_table(doc,
    ["服务", "镜像", "端口", "用途"],
    [
        ["postgres", "postgres:16-alpine", "5432", "关系型数据库"],
        ["elasticsearch", "elasticsearch:8.12.0", "9200", "全文搜索引擎"],
        ["milvus", "milvusdb/milvus:v2.3.7", "19530", "向量数据库"],
        ["etcd", "quay.io/coreos/etcd:v3.5.11", "2379", "Milvus元数据存储"],
        ["minio (应用)", "minio/minio", "9000/9001", "文件对象存储"],
        ["minio (Milvus)", "minio/minio", "9010/9011", "Milvus内部存储"],
        ["rabbitmq", "rabbitmq:3.13-management", "5672/15672", "消息队列"],
        ["redis", "redis:7-alpine", "6379", "缓存与任务队列"],
    ],
    col_widths=[3, 5, 3, 4]
)

p(doc, "所有服务通过 ai-literacy-net 桥接网络互联，数据通过 Docker Volume 持久化。")

# ── 27. 生产部署 ──
h2(doc, "27. 生产部署要点")
p(doc, "生产环境与开发环境的关键差异：", bold=True)

add_table(doc,
    ["配置项", "开发环境", "生产环境"],
    [
        ["DEBUG", "true", "false"],
        ["SECRET_KEY", "dev-secret-key...", "随机生成64位密钥"],
        ["JWT_SECRET_KEY", "dev-jwt-secret...", "随机生成64位密钥"],
        ["数据库密码", "ai_literacy_pass", "强随机密码"],
        ["MinIO密钥", "minioadmin", "强随机密码"],
        ["CORS", "localhost:3000", "实际域名"],
        ["API文档", "启用(/docs)", "禁用"],
        ["HTTPS", "否", "是（Let's Encrypt）"],
        ["反向代理", "无", "Nginx"],
    ],
    col_widths=[4, 5, 5]
)

p(doc, "Nginx配置要点：", bold=True)
bullet(doc, "反向代理 /api → FastAPI(8000)")
bullet(doc, "静态文件服务 → frontend/dist/")
bullet(doc, "SSL证书配置（Let's Encrypt/certbot）")
bullet(doc, "仅暴露80/443端口，其他服务内网访问")

p(doc, "密钥生成命令：", bold=True)
p(doc, '  python -c "import secrets; print(secrets.token_urlsafe(64))"')

# ── 28. 配置说明 ──
h2(doc, "28. 环境变量配置")
add_table(doc,
    ["变量名", "默认值", "说明"],
    [
        ["APP_NAME", "AI素养评测平台", "应用名称"],
        ["APP_VERSION", "0.1.0", "版本号"],
        ["DEBUG", "true", "调试模式（生产环境设为false）"],
        ["SECRET_KEY", "change-this", "应用密钥（务必修改）"],
        ["API_V1_PREFIX", "/api/v1", "API路径前缀"],
        ["POSTGRES_HOST", "localhost", "数据库主机"],
        ["POSTGRES_PORT", "5432", "数据库端口"],
        ["POSTGRES_USER", "ai_literacy", "数据库用户"],
        ["POSTGRES_PASSWORD", "ai_literacy_pass", "数据库密码"],
        ["POSTGRES_DB", "ai_literacy_db", "数据库名"],
        ["ELASTICSEARCH_HOST", "localhost", "搜索引擎主机"],
        ["ELASTICSEARCH_PORT", "9200", "搜索引擎端口"],
        ["MILVUS_HOST", "localhost", "向量数据库主机"],
        ["MILVUS_PORT", "19530", "向量数据库端口"],
        ["MINIO_HOST", "localhost", "对象存储主机"],
        ["MINIO_PORT", "9000", "对象存储端口"],
        ["MINIO_ACCESS_KEY", "minioadmin", "MinIO访问密钥"],
        ["MINIO_SECRET_KEY", "minioadmin", "MinIO密钥"],
        ["MINIO_BUCKET", "ai-literacy", "存储桶名称"],
        ["REDIS_HOST", "localhost", "Redis主机"],
        ["REDIS_PORT", "6379", "Redis端口"],
        ["RABBITMQ_HOST", "localhost", "RabbitMQ主机"],
        ["RABBITMQ_PORT", "5672", "RabbitMQ端口"],
        ["RABBITMQ_USER", "guest", "RabbitMQ用户"],
        ["RABBITMQ_PASSWORD", "guest", "RabbitMQ密码"],
        ["LLM_API_KEY", "your-api-key", "LLM API密钥"],
        ["LLM_BASE_URL", "https://api.deepseek.com/v1", "LLM API地址"],
        ["LLM_MODEL", "deepseek-reasoner", "LLM模型名"],
        ["JWT_SECRET_KEY", "change-this", "JWT签名密钥（务必修改）"],
        ["JWT_ALGORITHM", "HS256", "JWT算法"],
        ["JWT_ACCESS_TOKEN_EXPIRE_MINUTES", "30", "令牌有效期（分钟）"],
    ],
    col_widths=[5, 5, 5]
)

# ── 29. 数据库迁移 ──
h2(doc, "29. 数据库迁移")
p(doc, "系统使用 Alembic 管理数据库模式变更：")
bullet(doc, "初始化/更新：alembic upgrade head")
bullet(doc, "创建新迁移：alembic revision --autogenerate -m \"描述\"")
bullet(doc, "回退一版：alembic downgrade -1")
bullet(doc, "查看当前版本：alembic current")

p(doc, "当前迁移历史：", bold=True)
add_table(doc,
    ["版本号", "描述"],
    [
        ["25682f73fedf", "初始模式 - 全部核心表"],
        ["f6ab067da361", "用户软删除支持"],
        ["729a7ab81059", "排行榜可见性设置"],
        ["e08fd89ac436", "答题卡软删除支持"],
    ],
    col_widths=[5, 10]
)

# ── 30. 备份与恢复 ──
h2(doc, "30. 备份与恢复")

h3(doc, "30.1 数据库备份")
p(doc, "手动备份：")
p(doc, "  docker exec postgres pg_dump -U ai_literacy ai_literacy_db > backup.sql")
p(doc, "定时备份（每日凌晨2点）：")
p(doc, "  0 2 * * * docker exec postgres pg_dump -U ai_literacy ai_literacy_db > /opt/backups/db_$(date +%Y%m%d).sql")

h3(doc, "30.2 文件备份")
p(doc, "MinIO文件备份：")
p(doc, "  docker cp minio:/data ./minio_backup")

h3(doc, "30.3 恢复")
p(doc, "数据库恢复：")
p(doc, "  cat backup.sql | docker exec -i postgres psql -U ai_literacy ai_literacy_db")

# ── 31. FAQ ──
h2(doc, "31. 常见问题 (FAQ)")

p(doc, "Q1: 忘记管理员密码怎么办？", bold=True)
p(doc, "A: 可通过数据库直接重置。使用 bcrypt 生成新密码哈希后更新 users 表的 hashed_password 字段。")

p(doc, "Q2: LLM服务不可用时系统是否可以正常运行？", bold=True)
p(doc, "A: 可以。AI出题降级为模板生成，AI评分降级为规则匹配，其他核心功能不受影响。")

p(doc, "Q3: 如何修改默认管理员密码？", bold=True)
p(doc, "A: 登录后进入个人中心，使用「修改密码」功能。")

p(doc, "Q4: 批量导入用户失败怎么办？", bold=True)
p(doc, "A: 检查Excel/CSV文件是否包含username和email两列，确保格式正确且无重复用户名/邮箱。")

p(doc, "Q5: Docker容器无法启动怎么办？", bold=True)
p(doc, "A: 执行 docker compose logs 查看具体错误。常见原因：端口被占用、磁盘空间不足、内存不足。")

p(doc, "Q6: 如何切换LLM服务提供商？", bold=True)
p(doc, "A: 修改 .env 文件中的 LLM_BASE_URL、LLM_API_KEY、LLM_MODEL 三个变量，重启后端服务。")

p(doc, "Q7: 如何查看API文档？", bold=True)
p(doc, "A: 开发模式下访问 http://localhost:8000/docs（Swagger UI）或 /redoc（ReDoc）。")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
# 附录
# ═══════════════════════════════════════════════════════════════════
h1(doc, "附录")

# 附录A
h2(doc, "附录A  默认账号信息")
add_table(doc,
    ["服务", "用户名", "密码", "用途"],
    [
        ["平台管理员", "admin", "admin123", "系统管理"],
        ["PostgreSQL", "ai_literacy", "ai_literacy_pass", "数据库"],
        ["MinIO", "minioadmin", "minioadmin", "文件存储控制台(9001端口)"],
        ["RabbitMQ", "guest", "guest", "消息队列管理界面(15672端口)"],
    ],
    col_widths=[3, 3, 5, 4]
)
note_box(doc, "以上均为开发环境默认账号，生产环境务必全部修改！")

# 附录B
h2(doc, "附录B  端口列表")
add_table(doc,
    ["端口", "服务", "协议", "说明"],
    [
        ["3000", "Vue前端", "HTTP", "前端开发服务器"],
        ["8000", "FastAPI", "HTTP", "后端API服务"],
        ["1234", "LM Studio", "HTTP", "本地LLM服务"],
        ["5432", "PostgreSQL", "TCP", "关系型数据库"],
        ["9200", "Elasticsearch", "HTTP", "搜索引擎"],
        ["19530", "Milvus", "gRPC", "向量数据库"],
        ["9000", "MinIO API", "HTTP", "对象存储API"],
        ["9001", "MinIO Console", "HTTP", "存储管理界面"],
        ["5672", "RabbitMQ", "AMQP", "消息队列"],
        ["15672", "RabbitMQ Console", "HTTP", "队列管理界面"],
        ["6379", "Redis", "TCP", "缓存服务"],
        ["2379", "etcd", "HTTP", "Milvus元数据"],
    ],
    col_widths=[2, 4, 2, 6]
)

# 附录C
h2(doc, "附录C  支持的文件格式")
add_table(doc,
    ["格式分类", "扩展名", "MIME类型", "大小限制"],
    [
        ["PDF", ".pdf", "application/pdf", "100 MB"],
        ["Word", ".doc, .docx", "application/msword, ...openxml...", "50 MB"],
        ["Markdown", ".md, .markdown", "text/markdown", "10 MB"],
        ["HTML", ".html, .htm", "text/html", "10 MB"],
        ["图片", ".jpg/.jpeg/.png/.gif/.bmp/.webp", "image/*", "20 MB"],
        ["视频", ".mp4/.avi/.mov/.wmv/.flv/.mkv/.webm", "video/*", "500 MB"],
        ["音频", ".mp3/.wav/.flac/.aac/.ogg/.wma/.m4a", "audio/*", "200 MB"],
        ["CSV", ".csv", "text/csv", "50 MB"],
        ["JSON", ".json", "application/json", "50 MB"],
    ],
    col_widths=[3, 5, 4, 3]
)

# 附录D
h2(doc, "附录D  题型详细说明")
add_table(doc,
    ["题型", "标识", "选项结构", "正确答案格式", "评分方式"],
    [
        ["单选题", "single_choice", '{"A":"选项1","B":"选项2","C":"选项3","D":"选项4"}', "A", "自动（完全匹配）"],
        ["多选题", "multiple_choice", '{"A":"选项1","B":"选项2","C":"选项3","D":"选项4"}', "AB / ABC 等", "自动（完全匹配）"],
        ["判断题", "true_false", '{"A":"正确","B":"错误"}', "A 或 B", "自动（完全匹配）"],
        ["填空题", "fill_blank", "无选项", "文本答案", "自动+AI辅助"],
        ["简答题", "short_answer", "无选项", "参考答案", "AI多模型面板评分"],
    ],
    col_widths=[2, 3, 4, 3, 3]
)

# 附录E
h2(doc, "附录E  评分等级标准")
add_table(doc,
    ["等级", "中文", "得分率", "颜色", "证书"],
    [
        ["优秀", "Excellent", "≥ 90%", "金色 #FFD700", "可下载"],
        ["良好", "Good", "75% - 89%", "蓝色 #1890FF", "不可下载"],
        ["合格", "Pass", "60% - 74%", "绿色 #52C41A", "不可下载"],
        ["不合格", "Fail", "< 60%", "红色 #FF4D4F", "不可下载"],
    ],
    col_widths=[2, 3, 3, 4, 3]
)

p(doc, "")
p(doc, "五大AI素养评估维度：", bold=True)
add_table(doc,
    ["维度", "说明"],
    [
        ["AI基础知识", "人工智能基本概念、发展历史、核心技术原理"],
        ["AI技术应用", "AI工具使用能力、应用场景理解、实践操作技能"],
        ["AI伦理安全", "AI伦理意识、数据安全、隐私保护、社会影响认知"],
        ["AI批判思维", "对AI输出的评估能力、偏见识别、局限性理解"],
        ["AI创新实践", "利用AI解决问题、创新应用、跨学科融合能力"],
    ],
    col_widths=[4, 11]
)


# ═══════════════════════════════════════════════════════════════════
# 保存文档
# ═══════════════════════════════════════════════════════════════════
output_path = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "AI素养评测平台_系统说明书.docx"
)
doc.save(output_path)
print(f"文档已生成：{output_path}")
print(f"文件大小：{os.path.getsize(output_path) / 1024:.1f} KB")
