"""生成 AI素养评测平台 系统管理员手册 (Word .docx)"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# ── helpers ──────────────────────────────────────────────────────────
def set_cell_shading(cell, color_hex: str):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex, qn('w:val'): 'clear',
    })
    shading.append(shd)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for pg in cell.paragraphs:
            pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in pg.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '1F4E79')
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for pg in cell.paragraphs:
                for run in pg.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'E8F0FE')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def h1(doc, text):
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def h3(doc, text):
    doc.add_heading(text, level=3)


def h4(doc, text):
    doc.add_heading(text, level=4)


def p(doc, text, bold=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(10.5)
    return para


def bullet(doc, text, level=0):
    para = doc.add_paragraph(text, style='List Bullet')
    para.paragraph_format.left_indent = Cm(1.27 + level * 1.27)
    for run in para.runs:
        run.font.size = Pt(10)
    return para


def code_block(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    para.paragraph_format.left_indent = Cm(1.0)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    # 浅灰背景
    shading = para._element.get_or_add_pPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): 'F0F0F0', qn('w:val'): 'clear',
    })
    shading.append(shd)
    return para


def note_box(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(f"注意：{text}")
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x8B, 0x45, 0x13)
    run.italic = True
    return para


def tip_box(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(f"提示：{text}")
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x00, 0x6B, 0x3F)
    run.italic = True
    return para


# ── 开始构建文档 ─────────────────────────────────────────────────────
doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)


# ═══════════════════════════════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run("AI素养评测培训平台")
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("系统管理员手册")
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x2A, 0x6B, 0xA6)

doc.add_paragraph()

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run("vLLM 大模型集成 · 提示词配置 · 部署运维指南")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("版本：v0.1.0    更新日期：2026-03-08    适用角色：系统管理员")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
# 目录页
# ═══════════════════════════════════════════════════════════════════
h1(doc, "目录")
toc_items = [
    "第一章  系统概述",
    "第二章  系统架构与服务组件",
    "第三章  vLLM 大模型服务部署",
    "第四章  大模型调用与切换",
    "第五章  系统内置提示词一览",
    "第六章  可自定义提示词的入口",
    "第七章  平台部署与启动",
    "第八章  系统管理操作指南",
    "第九章  服务端口与访问地址",
    "第十章  日常运维与故障排查",
    "附录",
]
for item in toc_items:
    p(doc, item)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
# 第一章 系统概述
# ═══════════════════════════════════════════════════════════════════
h1(doc, "第一章  系统概述")

p(doc, "AI素养评测培训平台是一套面向企业和教育机构的AI素养评测与培训系统。平台围绕五大AI素养维度进行评测与培训。")

h2(doc, "1.1 五大AI素养维度")
add_table(doc,
    ["维度", "说明"],
    [
        ["AI基础知识", "机器学习原理、深度学习、算法基础、数据处理"],
        ["AI技术应用", "NLP、计算机视觉、推荐系统、大语言模型、AIGC"],
        ["AI伦理安全", "隐私保护、算法偏见、数据安全、负责任AI"],
        ["AI批判思维", "信息辨别、AI局限性认知、证据评估、逻辑推理"],
        ["AI创新实践", "提示工程、AI工具使用、工作流自动化、方案设计"],
    ],
    col_widths=[4, 12],
)

h2(doc, "1.2 核心功能模块")
bullet(doc, "智能题库管理：AI自动出题、多维度质量审核、题目校准与去重")
bullet(doc, "智能组卷与考试：自然语言组卷、自动组卷、在线考试、实时自动保存")
bullet(doc, "智能评分：客观题自动评分、主观题AI评分、多模型评委团评分（防偏见）")
bullet(doc, "情境化评测：多轮对话式情境判断测试（SJT）")
bullet(doc, "学习闭环：五维诊断报告、自适应学习路径、薄弱点训练")
bullet(doc, "运营管理：用户管理、组织管理、成绩导出、月度报告")

h2(doc, "1.3 用户角色")
add_table(doc,
    ["角色", "角色标识", "权限"],
    [
        ["管理员", "admin", "全部功能，含用户管理、系统配置"],
        ["组织者", "organizer", "素材上传、出题、组卷、考试管理"],
        ["审核员", "reviewer", "题目审核、答卷批阅"],
        ["考生", "examinee", "参加考试、查看成绩、学习训练"],
    ],
    col_widths=[3, 3, 10],
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
# 第二章 系统架构
# ═══════════════════════════════════════════════════════════════════
h1(doc, "第二章  系统架构与服务组件")

h2(doc, "2.1 系统架构图")
p(doc, "系统采用前后端分离架构，通过 8 个 AI 智能体（Agent）调用大模型服务：")

code_block(doc,
    "用户浏览器 → Nginx + Vue 3 前端 (:80)\n"
    "         → FastAPI 后端 (:8000)\n"
    "             → 8 个 AI 智能体 (Agent)\n"
    "                 → vLLM 大模型服务 (:8100)\n"
    "               Qwen/Qwen3.5-35B-A3B (GPTQ-Int4, MoE)"
)

h2(doc, "2.2 基础设施服务")
add_table(doc,
    ["服务", "镜像版本", "用途", "端口"],
    [
        ["PostgreSQL", "16-alpine", "关系型数据库", "5432"],
        ["Elasticsearch", "8.12.0", "全文搜索", "9200"],
        ["Milvus", "2.3.7", "向量数据库（语义检索/去重）", "19530"],
        ["MinIO", "RELEASE.2024-01", "对象存储（素材文件）", "9000/9001"],
        ["RabbitMQ", "3.13", "消息队列（异步任务）", "5672/15672"],
        ["Redis", "7-alpine", "缓存", "6379"],
        ["vLLM", "0.17.0", "大模型推理服务", "8100"],
    ],
    col_widths=[3, 4, 5.5, 3.5],
)

h2(doc, "2.3 AI 智能体清单")
add_table(doc,
    ["智能体", "文件", "功能"],
    [
        ["出题智能体", "question_agent.py", "根据知识内容生成多类型、多难度的评测题目"],
        ["评分智能体", "scoring_agent.py", "主观题AI评分，支持单模型和多评委团模式"],
        ["互动问答智能体", "interactive_agent.py", "多轮情境对话式情境判断测试（SJT）"],
        ["标注智能体", "annotation_agent.py", "自动标注教学材料的维度、难度、知识点"],
        ["审核智能体", "review_agent.py", "五维度AI质量审核题目"],
        ["意图解析智能体", "intent_agent.py", "将自然语言描述解析为组卷参数"],
        ["指标研究智能体", "indicator_agents.py", "AI趋势研究、指标建议、红队审核"],
        ["素材解析智能体", "material_agent.py", "素材内容提取与知识单元切分"],
    ],
    col_widths=[4, 4.5, 7.5],
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
# 第三章 vLLM 大模型服务部署
# ═══════════════════════════════════════════════════════════════════
h1(doc, "第三章  vLLM 大模型服务部署")

h2(doc, "3.1 环境要求")
add_table(doc,
    ["项目", "要求"],
    [
        ["操作系统", "Ubuntu 22.04+ / CentOS 8+ / 其他支持 CUDA 的 Linux"],
        ["GPU", "NVIDIA GPU，显存 >= 24GB（推荐 128GB+ 如 DGX Spark GB10）"],
        ["CUDA", "12.0+"],
        ["Python", "3.10+"],
    ],
    col_widths=[4, 12],
)

h2(doc, "3.2 安装 vLLM")
p(doc, "在服务器上创建独立的 Python 虚拟环境并安装 vLLM：")
code_block(doc, "python3 -m venv ~/vllm-env")
code_block(doc, "~/vllm-env/bin/pip install vllm")

h2(doc, "3.3 当前推荐模型")
p(doc, "平台当前使用 Qwen3.5-35B-A3B 模型（GPTQ-Int4 量化版）：")
add_table(doc,
    ["属性", "说明"],
    [
        ["模型全名", "Qwen/Qwen3.5-35B-A3B-GPTQ-Int4"],
        ["对外服务名", "Qwen/Qwen3.5-35B-A3B"],
        ["架构", "MoE — 总参数 35B，每次推理仅激活 3B"],
        ["量化", "GPTQ 4-bit（Marlin 后端）"],
        ["显存占用", "约 21 GiB（128GB 显存仅占 16%）"],
        ["KV Cache", "约 81 GiB 可用，支持约 106 万 tokens"],
        ["最大上下文", "262,144 tokens（256K）"],
        ["最大并发", "约 16x（基于 256K 上下文）"],
    ],
    col_widths=[4, 12],
)
note_box(doc, "为什么选择 GPTQ-Int4 而非 FP8？当前 vLLM 0.17.0 的 CUTLASS FP8 内核不兼容 NVIDIA Blackwell 架构（GB10, sm_121）。GPTQ-Int4 使用 Marlin 后端，完美兼容且显存占用更低。")

h2(doc, "3.4 下载并启动模型")

h3(doc, "方式一：直接启动（自动下载模型）")
p(doc, "首次启动会自动从 HuggingFace 镜像下载模型权重（约 23GB），后续启动直接使用缓存。")
code_block(doc,
    "# 设置 HuggingFace 镜像（国内服务器必须）\n"
    "export HF_ENDPOINT=https://hf-mirror.com\n"
    "\n"
    "# 启动 vLLM 服务（GPTQ-Int4 量化版）\n"
    "nohup ~/vllm-env/bin/python -m vllm.entrypoints.openai.api_server \\\n"
    "  --model Qwen/Qwen3.5-35B-A3B-GPTQ-Int4 \\\n"
    "  --served-model-name Qwen/Qwen3.5-35B-A3B \\\n"
    "  --host 0.0.0.0 \\\n"
    "  --port 8100 \\\n"
    "  --trust-remote-code \\\n"
    "  > ~/vllm-server.log 2>&1 &"
)

h3(doc, "方式二：先下载模型，再启动")
code_block(doc,
    "# 安装 huggingface-cli\n"
    "~/vllm-env/bin/pip install huggingface_hub\n"
    "\n"
    "# 使用镜像下载模型到指定目录\n"
    "export HF_ENDPOINT=https://hf-mirror.com\n"
    "~/vllm-env/bin/huggingface-cli download \\\n"
    "  --resume-download Qwen/Qwen3.5-35B-A3B-GPTQ-Int4 \\\n"
    "  --local-dir ~/models/Qwen3.5-35B-A3B-GPTQ-Int4\n"
    "\n"
    "# 从本地路径启动\n"
    "nohup ~/vllm-env/bin/python -m vllm.entrypoints.openai.api_server \\\n"
    "  --model ~/models/Qwen3.5-35B-A3B-GPTQ-Int4 \\\n"
    "  --served-model-name Qwen/Qwen3.5-35B-A3B \\\n"
    "  --host 0.0.0.0 --port 8100 \\\n"
    "  --trust-remote-code \\\n"
    "  > ~/vllm-server.log 2>&1 &"
)

h2(doc, "3.5 验证 vLLM 服务")
p(doc, "启动完成后（模型加载需要数分钟），使用以下命令验证：")
code_block(doc,
    "# 查看可用模型\n"
    "curl http://localhost:8100/v1/models\n"
    "\n"
    "# 测试对话能力\n"
    "curl http://localhost:8100/v1/chat/completions \\\n"
    "  -H 'Content-Type: application/json' \\\n"
    "  -d '{\n"
    '    "model": "Qwen/Qwen3.5-35B-A3B",\n'
    '    "messages": [{"role": "user", "content": "请介绍人工智能"}],\n'
    '    "max_tokens": 200\n'
    "  }'"
)

h2(doc, "3.6 vLLM 常用启动参数")
add_table(doc,
    ["参数", "说明", "推荐值"],
    [
        ["--model", "模型名称或本地路径", "Qwen/Qwen3.5-35B-A3B-GPTQ-Int4"],
        ["--served-model-name", "对外暴露的模型名称", "Qwen/Qwen3.5-35B-A3B"],
        ["--host", "监听地址", "0.0.0.0（允许外部访问）"],
        ["--port", "服务端口", "8100"],
        ["--trust-remote-code", "信任远程代码", "Qwen3.5 模型必须"],
        ["--tensor-parallel-size", "GPU张量并行数", "多卡时设为GPU数量"],
        ["--max-model-len", "最大序列长度", "默认262144（256K）"],
        ["--gpu-memory-utilization", "GPU显存利用率", "0.9（默认）"],
        ["--dtype", "数据精度", "auto（GPTQ自动使用bfloat16）"],
        ["--enforce-eager", "禁用torch.compile", "GPU兼容性问题时使用"],
    ],
    col_widths=[5, 5.5, 5.5],
)
note_box(doc, "--served-model-name 非常重要。它决定了 API 对外暴露的模型名称，必须与 .env 中的 LLM_MODEL 值一致。")

h2(doc, "3.7 设置 vLLM 开机自启（systemd）")
p(doc, "创建服务文件 /etc/systemd/system/vllm.service：")
code_block(doc,
    "[Unit]\n"
    "Description=vLLM OpenAI-compatible API Server\n"
    "After=network.target\n"
    "\n"
    "[Service]\n"
    "Type=simple\n"
    "User=dell\n"
    'Environment="HF_ENDPOINT=https://hf-mirror.com"\n'
    "ExecStart=/home/dell/vllm-env/bin/python \\\n"
    "  -m vllm.entrypoints.openai.api_server \\\n"
    "  --model Qwen/Qwen3.5-35B-A3B-GPTQ-Int4 \\\n"
    "  --served-model-name Qwen/Qwen3.5-35B-A3B \\\n"
    "  --host 0.0.0.0 --port 8100 \\\n"
    "  --trust-remote-code\n"
    "Restart=on-failure\n"
    "RestartSec=10\n"
    "\n"
    "[Install]\n"
    "WantedBy=multi-user.target"
)
p(doc, "启用服务：")
code_block(doc,
    "sudo systemctl daemon-reload\n"
    "sudo systemctl enable vllm\n"
    "sudo systemctl start vllm\n"
    "sudo systemctl status vllm      # 查看状态\n"
    "journalctl -u vllm -f           # 查看日志"
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
# 第四章 大模型调用与切换
# ═══════════════════════════════════════════════════════════════════
h1(doc, "第四章  大模型调用与切换")

h2(doc, "4.1 配置文件说明")
p(doc, "平台通过 3 个环境变量控制大模型连接，定义在 .env 或 .env.production 中：")
add_table(doc,
    ["变量", "说明", "当前值"],
    [
        ["LLM_API_KEY", "API 密钥", "token-not-needed（本地vLLM不需要）"],
        ["LLM_BASE_URL", "API 地址", "http://localhost:8100/v1"],
        ["LLM_MODEL", "模型标识", "Qwen/Qwen3.5-35B-A3B"],
    ],
    col_widths=[4, 4, 8],
)

p(doc, "配置文件位置：")
add_table(doc,
    ["文件", "用途"],
    [
        [".env", "开发环境配置"],
        [".env.production", "生产环境配置（Docker部署使用）"],
        ["app/core/config.py", "默认值定义（.env 未设置时的降级值）"],
    ],
    col_widths=[5, 11],
)

h2(doc, "4.2 切换大模型的方法")

h3(doc, "场景一：切换 vLLM 加载的模型")
p(doc, "1. 停止当前 vLLM 服务：")
code_block(doc, 'pkill -f "vllm.entrypoints.openai.api_server"')
p(doc, "2. 用新模型启动 vLLM：")
code_block(doc,
    "export HF_ENDPOINT=https://hf-mirror.com\n"
    "nohup ~/vllm-env/bin/python -m vllm.entrypoints.openai.api_server \\\n"
    "  --model <新模型HF名称> \\\n"
    "  --served-model-name <对外服务名> \\\n"
    "  --host 0.0.0.0 --port 8100 \\\n"
    "  --trust-remote-code \\\n"
    "  > ~/vllm-server.log 2>&1 &"
)
p(doc, "3. 更新 .env 中的 LLM_MODEL（必须与 --served-model-name 一致）")
p(doc, "4. 重启后端服务：docker compose restart app")

h3(doc, "场景二：切换到云端 API")
p(doc, "修改 .env 或 .env.production：")
code_block(doc,
    "# DeepSeek API\n"
    "LLM_API_KEY=sk-xxxxxxxxxxxxxxxx\n"
    "LLM_BASE_URL=https://api.deepseek.com/v1\n"
    "LLM_MODEL=deepseek-chat\n"
    "\n"
    "# OpenAI API\n"
    "LLM_API_KEY=sk-xxxxxxxxxxxxxxxx\n"
    "LLM_BASE_URL=https://api.openai.com/v1\n"
    "LLM_MODEL=gpt-4o\n"
    "\n"
    "# 其他 OpenAI 兼容 API（如硅基流动、智谱 AI 等）\n"
    "LLM_API_KEY=你的API密钥\n"
    "LLM_BASE_URL=https://api.provider.com/v1\n"
    "LLM_MODEL=provider-model-name"
)
p(doc, "修改后重启后端：docker compose restart app")

h3(doc, "场景三：切换回本地 vLLM")
code_block(doc,
    "LLM_API_KEY=token-not-needed\n"
    "LLM_BASE_URL=http://localhost:8100/v1\n"
    "LLM_MODEL=Qwen/Qwen3.5-35B-A3B"
)
note_box(doc, "Docker 容器内访问宿主机上的 vLLM 时，.env.production 中应使用 http://host.docker.internal:8100/v1，而非 localhost。")

h2(doc, "4.3 推荐模型列表")
h3(doc, "本地部署模型（vLLM）")
add_table(doc,
    ["模型", "架构", "显存需求", "推荐场景"],
    [
        ["Qwen/Qwen3.5-35B-A3B-GPTQ-Int4", "MoE 35B/3B", "21GB", "⭐ 当前使用，Blackwell兼容"],
        ["Qwen/Qwen3.5-35B-A3B", "MoE 35B/3B", "70GB", "BF16全精度版"],
        ["Qwen/Qwen2.5-7B-Instruct", "Dense 7B", "16GB", "轻量部署，入门选择"],
        ["Qwen/Qwen2.5-14B-Instruct", "Dense 14B", "32GB", "中等质量需求"],
        ["Qwen/Qwen2.5-72B-Instruct-AWQ", "Dense 72B", "48GB", "高质量（AWQ量化）"],
    ],
    col_widths=[5.5, 3, 2.5, 5],
)
tip_box(doc, 'MoE 架构说明：Qwen3.5-35B-A3B 总参数 35B，但每次推理仅激活 3B 参数，实现\u201c大模型质量 + 小模型速度\u201d的最佳平衡。')

h3(doc, "Blackwell GPU (GB10) 兼容性")
add_table(doc,
    ["量化格式", "vLLM 兼容性", "使用内核", "说明"],
    [
        ["GPTQ-Int4", "✅ 兼容", "Marlin", "推荐，当前使用"],
        ["BF16（无量化）", "✅ 兼容", "FlashAttention v2", "显存占用大（约70GB）"],
        ["FP8", "❌ 不兼容", "CUTLASS", "vLLM 0.17.0 不支持 sm_121"],
        ["AWQ", "✅ 兼容", "Marlin", "可作为备选方案"],
    ],
    col_widths=[3.5, 3, 3.5, 6],
)
note_box(doc, "在 Blackwell GPU（DGX Spark GB10）上，请避免使用 FP8 量化模型。待 vLLM 后续版本更新后可重新评估。")

h3(doc, "云端 API 模型")
add_table(doc,
    ["模型", "提供商", "说明"],
    [
        ["deepseek-chat", "DeepSeek", "性价比高，中文优秀"],
        ["gpt-4o", "OpenAI", "综合能力最强"],
        ["glm-4", "智谱 AI", "国产大模型"],
    ],
    col_widths=[4, 4, 8],
)

h2(doc, "4.4 智能降级机制")
p(doc, "当大模型服务不可用时，系统自动降级为规则/模板模式，保证平台基本功能可用：")
add_table(doc,
    ["功能", "LLM 可用时", "降级模式"],
    [
        ["出题", "AI智能生成高质量多样化题目", "基于模板生成（质量有限）"],
        ["评分", "AI语义理解评分", "关键词匹配 + 规则评分"],
        ["互动问答", "多轮情境对话", "预设回复模板"],
        ["审核", "AI五维度质量评估", "基础格式检查"],
        ["标注", "AI自动内容标注", "关键词匹配分类"],
        ["意图解析", "自然语言理解组卷", "返回错误提示"],
    ],
    col_widths=[3, 6, 7],
)
p(doc, "降级触发条件：")
bullet(doc, "LLM_API_KEY 为默认值 your-api-key 时直接降级")
bullet(doc, "API 调用超时（60秒）或返回错误时降级")
bullet(doc, "降级不影响平台其他功能正常使用")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
# 第五章 系统内置提示词一览
# ═══════════════════════════════════════════════════════════════════
h1(doc, "第五章  系统内置提示词一览")

p(doc, "平台共内置 11 个系统提示词（System Prompt），分布在 8 个 AI 智能体中。以下列出每个提示词的用途、位置和核心内容。")

h2(doc, "5.1 出题智能体")
p(doc, "文件位置：app/agents/question_agent.py，第 215-294 行", bold=True)
p(doc, "角色定义：资深AI素养评测出题专家，10年以上教育测评经验")
p(doc, "核心指令：")
bullet(doc, "布鲁姆认知目标分类法（Bloom's Taxonomy）精准应用")
bullet(doc, "干扰项心理学（Distractor Psychology）设计高质量干扰项")
bullet(doc, "情境丰富性：优先使用具体场景包裹知识点")
bullet(doc, "10种题目风格混合使用：直接知识型、情景应用型、案例分析型、对比辨析型、问题解决型、因果推理型、伦理困境型、评价判断型、趋势预测型、实践操作型")
bullet(doc, "5级难度校准（入门/简单/中等/困难/专家）")
bullet(doc, "正确答案位置分散要求")
p(doc, "输出格式：JSON数组，每道题包含 question_type、stem、options、correct_answer、explanation、knowledge_tags、dimension")

h2(doc, "5.2 评分智能体")
p(doc, "文件位置：app/agents/scoring_agent.py", bold=True)

h3(doc, "5.2.1 单模型评分提示词（第 16-34 行）")
p(doc, "角色定义：专业AI素养评测评分专家")
p(doc, "评分规则：根据参考答案覆盖程度给分；关键概念必须准确；表述清晰可加分；明显错误应扣分")
p(doc, "输出格式：{\"earned_ratio\": 0.0-1.0, \"feedback\": \"评分反馈\"}")

h3(doc, "5.2.2 多模型评委团提示词（第 157-185 行）")
p(doc, "防偏见机制：")
bullet(doc, "不因回答长度给额外分数")
bullet(doc, "不因礼貌/讨好语言加分")
bullet(doc, "只关注内容准确性、完整性和逻辑性")
bullet(doc, "位置交换策略（先读参考答案/先读学生作答/同时对比）")
p(doc, "四维度评分：accuracy（准确性）、completeness（完整性）、logic（逻辑性）、expression（表达性）")

h2(doc, "5.3 互动问答智能体")
p(doc, "文件位置：app/agents/interactive_agent.py", bold=True)

h3(doc, "5.3.1 情境对话提示词（第 17-57 行）")
p(doc, "角色定义：AI素养情境评测专家，主持情境判断测试（SJT）")
p(doc, "动态参数（由系统注入）：")
bullet(doc, "{role_description}：角色描述（如\"AI产品经理\"）")
bullet(doc, "{scenario}：场景背景")
bullet(doc, "{dimension}：评估维度")
bullet(doc, "{difficulty}：难度等级 1-5")
p(doc, "评估三维度：prompt_engineering（提示工程能力）、critical_thinking（批判性思维）、ethical_decision（伦理决策）")

h3(doc, "5.3.2 会话总结提示词（第 60-82 行）")
p(doc, "根据多轮对话记录生成最终评估摘要，包含总分、维度评分、关键决策点、优势、不足、建议。")

h2(doc, "5.4 标注智能体")
p(doc, "文件位置：app/agents/annotation_agent.py，第 16-33 行", bold=True)
p(doc, "功能：分析教学材料内容，自动标注维度、难度、知识点、摘要、标签")

h2(doc, "5.5 审核智能体")
p(doc, "文件位置：app/agents/review_agent.py，第 20-46 行", bold=True)
p(doc, "五维度审核标准：")
add_table(doc,
    ["维度", "评分范围", "审核内容"],
    [
        ["stem_clarity（题干清晰度）", "1-5分", "表述是否清晰、无歧义"],
        ["option_quality（选项质量）", "1-5分", "干扰项是否合理"],
        ["answer_correctness（答案正确性）", "1-5分", "正确答案是否确实正确"],
        ["knowledge_alignment（知识对齐）", "1-5分", "与标注维度是否一致"],
        ["difficulty_calibration（难度校准）", "1-5分", "难度标注是否匹配"],
    ],
    col_widths=[5.5, 2.5, 8],
)
p(doc, "审核决策：approve（通过，>=3.5）、revise（修订，2.5-3.5）、reject（拒绝，<2.5）")

h2(doc, "5.6 意图解析智能体")
p(doc, "文件位置：app/agents/intent_agent.py，第 42-70 行", bold=True)
p(doc, "功能：将自然语言描述解析为组卷参数")
p(doc, "示例输入：\"出一套20道的AI入门测试，包含15道单选和5道判断，限时30分钟\"")

h2(doc, "5.7 指标体系智能体")
p(doc, "文件位置：app/agents/indicator_agents.py", bold=True)
p(doc, "包含 3 个子智能体：")
add_table(doc,
    ["智能体", "行号", "功能", "temperature"],
    [
        ["研究员", "37-59", "分析AI发展趋势", "0.7"],
        ["顾问", "89-107", "提出评测指标更新建议", "0.5"],
        ["红队审核员", "136-158", "审查指标提案的可行性", "0.3"],
    ],
    col_widths=[3, 2.5, 6.5, 4],
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
# 第六章 可自定义提示词的入口
# ═══════════════════════════════════════════════════════════════════
h1(doc, "第六章  可自定义提示词的入口")

p(doc, "系统提供了多个入口允许管理员和使用者输入自定义提示词，以影响AI的输出。")

h2(doc, "6.1 题目生成 — 自定义出题要求")
p(doc, "API 入口：")
add_table(doc,
    ["端点", "方法", "说明"],
    [
        ["/api/v1/questions/generate", "POST", "基于知识单元生成"],
        ["/api/v1/questions/generate/material/{id}", "POST", "基于素材批量生成"],
        ["/api/v1/questions/generate/free", "POST", "自由生成（必须提供custom_prompt）"],
        ["/api/v1/questions/generate/bank/{id}", "POST", "构建题库"],
    ],
    col_widths=[6.5, 2, 7.5],
)

p(doc, "自定义提示词参数：custom_prompt（字符串）", bold=True)
p(doc, "作用方式：系统将 custom_prompt 作为「额外要求」追加到出题指令末尾")

p(doc, "可控参数一览：")
add_table(doc,
    ["参数", "类型", "说明"],
    [
        ["question_types", "列表", "single_choice, multiple_choice, true_false, fill_blank, short_answer"],
        ["count", "整数", "生成题目数量"],
        ["difficulty", "1-5", "难度等级"],
        ["bloom_level", "字符串", "布鲁姆认知层次：remember/understand/apply/analyze/evaluate/create"],
        ["custom_prompt", "字符串", "自定义提示词（追加到系统提示词后）"],
    ],
    col_widths=[4, 2.5, 9.5],
)

p(doc, "示例请求体：")
code_block(doc,
    '{\n'
    '  "content": "关于大语言模型的知识内容...",\n'
    '  "question_types": ["single_choice", "multiple_choice"],\n'
    '  "count": 5,\n'
    '  "difficulty": 3,\n'
    '  "bloom_level": "apply",\n'
    '  "custom_prompt": "请侧重考查提示工程方面的知识，\n'
    '    题目场景以办公自动化为主，避免涉及编程代码"\n'
    '}'
)

h2(doc, "6.2 主观题评分 — 自定义评分标准（Rubric）")
p(doc, "API 入口：POST /api/v1/scores/grade/{sheet_id}", bold=True)
p(doc, "自定义参数：评分时可传入 rubric 字典作为自定义评分标准")
p(doc, "作用方式：系统将 rubric 插入评分提示词中作为「评分标准（Rubric）」")
code_block(doc,
    '{\n'
    '  "rubric": {\n'
    '    "概念准确性": "准确使用AI相关术语，定义无误（40%）",\n'
    '    "论述完整性": "覆盖所有要求的知识点（30%）",\n'
    '    "实例运用": "能举出恰当的实际案例（20%）",\n'
    '    "表达规范": "语言清晰、逻辑连贯（10%）"\n'
    '  }\n'
    '}'
)

h2(doc, "6.3 多模型评委团评分")
p(doc, "API 入口：POST /api/v1/scores/panel-score", bold=True)
add_table(doc,
    ["参数", "说明"],
    [
        ["num_evaluators", "评委数量（2-5，默认3）"],
        ["rubric", "自定义评分标准字典"],
    ],
    col_widths=[4, 12],
)
p(doc, "工作原理：系统使用同一个LLM模型模拟多个独立评委，通过位置交换策略消除偏见，最终取中位数得分。")

h2(doc, "6.4 互动问答 — 场景与角色定义")
p(doc, "API 入口：POST /api/v1/interactive", bold=True)
add_table(doc,
    ["参数", "说明", "示例"],
    [
        ["scenario", "场景描述", "你是一家初创公司的CTO，正在考虑引入AI功能"],
        ["role_description", "AI扮演的角色", "资深AI产品顾问"],
        ["dimension", "评估维度", "AI技术应用"],
        ["difficulty", "初始难度", "3"],
    ],
    col_widths=[4, 4, 8],
)
p(doc, "这些参数会被注入到情境对话的系统提示词模板中。")

h2(doc, "6.5 自然语言组卷")
p(doc, "API 入口：POST /api/v1/exams/intent/assemble", bold=True)
p(doc, "自定义参数：description（自然语言描述）")
p(doc, "示例：")
code_block(doc,
    '{\n'
    '  "description": "为新入职的产品经理设计一套AI素养评测，\n'
    '    20道题，以应用和伦理为重点，难度中等偏上，时间30分钟"\n'
    '}'
)
p(doc, "系统通过意图解析智能体将自然语言转化为结构化组卷参数并自动组卷。")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
# 第七章 平台部署与启动
# ═══════════════════════════════════════════════════════════════════
h1(doc, "第七章  平台部署与启动")

h2(doc, "7.1 前置条件")
bullet(doc, "Docker 和 Docker Compose 已安装")
bullet(doc, "vLLM 服务已启动（参见第三章）")
bullet(doc, "GPU 驱动已正确安装")

h2(doc, "7.2 一键部署")
code_block(doc,
    "cd ~/ai-literacy-platform\n"
    "\n"
    "# 检查/编辑生产配置（必须修改标注为 [必改] 的项目）\n"
    "vim .env.production\n"
    "\n"
    "# 运行部署脚本\n"
    "chmod +x deploy.sh\n"
    "./deploy.sh"
)
p(doc, "部署脚本自动完成：")
bullet(doc, "1. 检查 Docker 环境")
bullet(doc, "2. 检查 vLLM 大模型服务是否就绪")
bullet(doc, "3. 验证配置文件，自动生成安全密钥")
bullet(doc, "4. 构建 Docker 镜像")
bullet(doc, "5. 启动所有服务")
bullet(doc, "6. 等待就绪并输出访问地址")

h2(doc, "7.3 手动部署")
code_block(doc,
    "# 构建镜像\n"
    "docker compose build\n"
    "\n"
    "# 启动所有服务\n"
    "docker compose up -d\n"
    "\n"
    "# 查看日志\n"
    "docker compose logs -f app\n"
    "\n"
    "# 运行数据库迁移\n"
    "docker compose exec app alembic upgrade head"
)

h2(doc, "7.4 生产环境配置要点")
p(doc, ".env.production 中必须修改的配置项：")
add_table(doc,
    ["配置项", "说明"],
    [
        ["SECRET_KEY", "应用密钥（使用 openssl rand -hex 32 生成）"],
        ["JWT_SECRET_KEY", "JWT 密钥（使用 openssl rand -hex 32 生成）"],
        ["POSTGRES_PASSWORD", "数据库密码（建议强密码）"],
        ["LLM_BASE_URL", "Docker 内使用 http://host.docker.internal:8100/v1"],
    ],
    col_widths=[5, 11],
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
# 第八章 系统管理操作指南
# ═══════════════════════════════════════════════════════════════════
h1(doc, "第八章  系统管理操作指南")

h2(doc, "8.1 用户管理")
add_table(doc,
    ["操作", "API", "说明"],
    [
        ["创建用户", "POST /api/v1/users", "指定 username, email, password, role"],
        ["批量导入", "POST /api/v1/users/batch-import", "上传 Excel/CSV 文件"],
        ["重置密码", "POST /api/v1/users/{id}/reset-password", "管理员重置用户密码"],
        ["停用用户", "DELETE /api/v1/users/{id}", "软删除（可恢复）"],
        ["恢复用户", "POST /api/v1/users/{id}/restore", "恢复已删除用户"],
        ["查看列表", "GET /api/v1/users", "支持 role/keyword/status 筛选"],
    ],
    col_widths=[3, 6, 7],
)

h2(doc, "8.2 题库管理")
add_table(doc,
    ["操作", "API"],
    [
        ["AI 批量出题", "POST /api/v1/questions/generate/material/{id}"],
        ["AI 自由出题", "POST /api/v1/questions/generate/free"],
        ["AI 质量检查", "POST /api/v1/questions/{id}/ai-check"],
        ["批量审核", "POST /api/v1/questions/batch/review"],
        ["自动标记低质量题", "POST /api/v1/questions/calibration/auto-flag"],
        ["查找重复题", "GET /api/v1/questions/calibration/similar"],
        ["全局质量报告", "GET /api/v1/questions/analysis/report"],
        ["导出题库 (Markdown)", "POST /api/v1/questions/batch/export-md"],
        ["导入题库 (Markdown)", "POST /api/v1/questions/batch/import-md"],
    ],
    col_widths=[5, 11],
)

h2(doc, "8.3 考试管理")
add_table(doc,
    ["操作", "API"],
    [
        ["创建考试", "POST /api/v1/exams"],
        ["自然语言组卷", "POST /api/v1/exams/intent/assemble"],
        ["智能自动组卷", "POST /api/v1/exams/{id}/assemble/auto"],
        ["发布考试", "POST /api/v1/exams/{id}/publish"],
        ["关闭考试", "POST /api/v1/exams/{id}/close"],
        ["考试数据分析", "GET /api/v1/exams/{id}/analysis"],
    ],
    col_widths=[5, 11],
)

h2(doc, "8.4 成绩管理")
add_table(doc,
    ["操作", "API"],
    [
        ["查看所有成绩", "GET /api/v1/scores/all"],
        ["导出成绩 Excel", "POST /api/v1/scores/export"],
        ["批阅主观题", "POST /api/v1/scores/grade/{sheet_id}"],
        ["评委团评分", "POST /api/v1/scores/panel-score"],
        ["五维诊断报告", "GET /api/v1/scores/{id}/diagnostic"],
        ["排行榜", "GET /api/v1/scores/leaderboard"],
    ],
    col_widths=[5, 11],
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
# 第九章 服务端口与访问地址
# ═══════════════════════════════════════════════════════════════════
h1(doc, "第九章  服务端口与访问地址")

h2(doc, "9.1 端口一览")
add_table(doc,
    ["端口", "服务", "说明"],
    [
        ["80", "Nginx + 前端", "平台访问入口"],
        ["8000", "FastAPI 后端", "API 接口 + Swagger 文档"],
        ["8100", "vLLM", "大模型推理服务"],
        ["5432", "PostgreSQL", "数据库"],
        ["9200", "Elasticsearch", "全文搜索"],
        ["19530", "Milvus", "向量数据库"],
        ["9000", "MinIO API", "对象存储"],
        ["9001", "MinIO Console", "对象存储管理界面"],
        ["5672", "RabbitMQ", "消息队列"],
        ["15672", "RabbitMQ Management", "消息队列管理界面"],
        ["6379", "Redis", "缓存"],
    ],
    col_widths=[3, 5, 8],
)

h2(doc, "9.2 常用访问地址")
add_table(doc,
    ["用途", "地址"],
    [
        ["平台首页", "http://<服务器IP>"],
        ["API 文档 (Swagger)", "http://<服务器IP>:8000/docs"],
        ["vLLM 模型列表", "http://<服务器IP>:8100/v1/models"],
        ["MinIO 管理", "http://<服务器IP>:9001（minioadmin/minioadmin）"],
        ["RabbitMQ 管理", "http://<服务器IP>:15672（guest/guest）"],
    ],
    col_widths=[5, 11],
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
# 第十章 日常运维与故障排查
# ═══════════════════════════════════════════════════════════════════
h1(doc, "第十章  日常运维与故障排查")

h2(doc, "10.1 常用运维命令")
code_block(doc,
    "# 查看所有服务状态\n"
    "docker compose ps\n"
    "\n"
    "# 查看后端日志\n"
    "docker compose logs -f app\n"
    "\n"
    "# 重启后端（配置变更后）\n"
    "docker compose restart app\n"
    "\n"
    "# 重启所有服务\n"
    "docker compose restart\n"
    "\n"
    "# 停止所有服务\n"
    "docker compose down\n"
    "\n"
    "# 停止并清除数据卷（注意：数据全部丢失）\n"
    "docker compose down -v"
)

h2(doc, "10.2 vLLM 运维")
code_block(doc,
    "# 查看 vLLM 日志\n"
    "tail -f ~/vllm-server.log\n"
    "\n"
    "# 检查 vLLM 进程\n"
    "ps aux | grep vllm\n"
    "\n"
    "# 查看 GPU 使用情况\n"
    "nvidia-smi\n"
    "\n"
    "# 停止 vLLM\n"
    'pkill -f "vllm.entrypoints.openai.api_server"\n'
    "\n"
    "# 重启 vLLM（当前模型：Qwen3.5-35B-A3B GPTQ-Int4）\n"
    "export HF_ENDPOINT=https://hf-mirror.com\n"
    "nohup ~/vllm-env/bin/python -m vllm.entrypoints.openai.api_server \\\n"
    "  --model Qwen/Qwen3.5-35B-A3B-GPTQ-Int4 \\\n"
    "  --served-model-name Qwen/Qwen3.5-35B-A3B \\\n"
    "  --host 0.0.0.0 --port 8100 \\\n"
    "  --trust-remote-code \\\n"
    "  > ~/vllm-server.log 2>&1 &\n"
    "\n"
    "# 查看 vLLM 当前加载的模型\n"
    "curl -s http://localhost:8100/v1/models | python3 -m json.tool"
)

h2(doc, "10.3 常见问题排查")

h3(doc, "Q1：AI 出题/评分返回模板化结果（质量不佳）")
p(doc, "原因：大模型服务不可用，系统自动降级为模板模式")
p(doc, "排查步骤：")
bullet(doc, "检查 vLLM 是否运行：curl http://localhost:8100/v1/models")
bullet(doc, "检查后端日志：docker compose logs app | grep LLM")
bullet(doc, "确认 .env.production 中 LLM_BASE_URL 配置正确")
bullet(doc, "Docker 内检查：docker compose exec app curl http://host.docker.internal:8100/v1/models")

h3(doc, "Q2：vLLM 启动后 GPU 内存不足 (OOM)")
p(doc, "解决方案：")
bullet(doc, "方案1：降低显存利用率 --gpu-memory-utilization 0.7")
bullet(doc, "方案2：减小最大序列长度 --max-model-len 4096")
bullet(doc, "方案3：使用 GPTQ-Int4 量化模型（当前已使用，仅 21GB）")
bullet(doc, "方案4：使用更小的模型 --model Qwen/Qwen2.5-7B-Instruct")
note_box(doc, "在 Blackwell GPU (GB10) 上不要使用 FP8 量化模型，会因 CUTLASS 内核不兼容而启动失败。")

h3(doc, "Q3：Docker 容器无法连接 vLLM")
p(doc, "原因：docker-compose.yml 中需配置 extra_hosts")
p(doc, "确认以下配置存在：")
code_block(doc,
    "app:\n"
    "  extra_hosts:\n"
    '    - "host.docker.internal:host-gateway"\n'
    "  environment:\n"
    "    - LLM_BASE_URL=http://host.docker.internal:8100/v1"
)

h3(doc, "Q4：模型下载超时或失败")
code_block(doc,
    "# 设置 HuggingFace 镜像（国内必须）\n"
    "export HF_ENDPOINT=https://hf-mirror.com\n"
    "\n"
    "# 使用 huggingface-cli 下载（支持断点续传）\n"
    "~/vllm-env/bin/huggingface-cli download \\\n"
    "  --resume-download Qwen/Qwen3.5-35B-A3B-GPTQ-Int4\n"
    "\n"
    "# 检查下载进度\n"
    "du -sh ~/.cache/huggingface/hub/models--Qwen--*\n"
    "\n"
    "# 清理不完整的下载\n"
    'find ~/.cache/huggingface/hub/ -name "*.incomplete" -delete'
)

h3(doc, "Q5：如何确认大模型的实际调用效果")
code_block(doc,
    "curl http://localhost:8100/v1/chat/completions \\\n"
    "  -H 'Content-Type: application/json' \\\n"
    "  -d '{\n"
    '    "model": "Qwen/Qwen3.5-35B-A3B",\n'
    '    "messages": [\n'
    '      {"role": "system", "content": "你是AI素养评测出题专家"},\n'
    '      {"role": "user", "content": "请生成1道关于机器学习的单选题"}\n'
    "    ],\n"
    '    "temperature": 0.7,\n'
    '    "max_tokens": 500\n'
    "  }'"
)
tip_box(doc, 'Qwen3.5 模型默认启用\u201c思考链\u201d模式，会在回答前输出推理过程。平台的各 Agent 已能正确处理此行为。')

h3(doc, "Q6：vLLM 启动时报 CUTLASS / cutlass_scaled_mm 错误")
p(doc, "原因：FP8 量化模型在 Blackwell GPU (sm_121) 上不兼容")
p(doc, "解决方案：")
bullet(doc, "改用 GPTQ-Int4 量化版本：--model Qwen/Qwen3.5-35B-A3B-GPTQ-Int4")
bullet(doc, "或使用 BF16 全精度：--model Qwen/Qwen3.5-35B-A3B --dtype bfloat16 --enforce-eager")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
# 附录
# ═══════════════════════════════════════════════════════════════════
h1(doc, "附录")

h2(doc, "附录A  各智能体 LLM 调用参数汇总")
add_table(doc,
    ["智能体", "文件", "temperature", "max_tokens", "用途"],
    [
        ["出题", "question_agent.py", "0.7", "4096", "生成多样化题目"],
        ["评分（单模型）", "scoring_agent.py", "0.2", "512", "精确评分"],
        ["评分（评委团）", "scoring_agent.py", "0.3", "512", "多角度评分"],
        ["互动对话", "interactive_agent.py", "0.7", "800", "情境对话"],
        ["互动总结", "interactive_agent.py", "0.3", "800", "生成摘要"],
        ["标注", "annotation_agent.py", "0.3", "500", "内容分析"],
        ["审核", "review_agent.py", "0.3", "1024", "质量检查"],
        ["意图解析", "intent_agent.py", "0.3", "500", "自然语言理解"],
        ["研究员", "indicator_agents.py", "0.7", "1000", "趋势分析"],
        ["顾问", "indicator_agents.py", "0.5", "1000", "指标建议"],
        ["红队审核", "indicator_agents.py", "0.3", "1000", "提案审查"],
    ],
    col_widths=[3.5, 4, 2.5, 2.5, 3.5],
)

h2(doc, "附录B  题型与布鲁姆层次对照")
add_table(doc,
    ["题型", "适用布鲁姆层次", "说明"],
    [
        ["单选题 (single_choice)", "全部层次", "最通用题型"],
        ["多选题 (multiple_choice)", "理解、分析、评价", "考查综合判断"],
        ["判断题 (true_false)", "记忆、理解", "概念辨析"],
        ["填空题 (fill_blank)", "记忆、理解", "术语掌握"],
        ["简答题 (short_answer)", "理解、应用、分析", "深度考查"],
    ],
    col_widths=[5, 5, 6],
)

h2(doc, "附录C  五维度关键词映射")
p(doc, "系统使用关键词匹配自动分类题目所属维度：")
add_table(doc,
    ["维度", "部分关键词"],
    [
        ["AI基础知识", "机器学习、深度学习、神经网络、算法、Transformer、梯度下降、反向传播"],
        ["AI技术应用", "NLP、计算机视觉、语音识别、ChatGPT、大语言模型、AIGC、推荐系统"],
        ["AI伦理安全", "隐私、偏见、公平、deepfake、数据保护、算法歧视、可解释性"],
        ["AI批判思维", "批判、局限、评估、验证、信息素养、逻辑、谬误"],
        ["AI创新实践", "提示工程、prompt、AI工具、微调、API、部署、自动化"],
    ],
    col_widths=[3.5, 12.5],
)

h2(doc, "附录D  配置文件环境变量完整清单")
add_table(doc,
    ["变量", "默认值", "说明"],
    [
        ["APP_NAME", "AI素养评测平台", "应用名称"],
        ["APP_VERSION", "0.1.0", "版本号"],
        ["DEBUG", "false", "调试模式（生产环境设为false）"],
        ["SECRET_KEY", "[必改]", "应用密钥"],
        ["API_V1_PREFIX", "/api/v1", "API 路径前缀"],
        ["POSTGRES_HOST", "postgres", "数据库地址（Docker服务名）"],
        ["POSTGRES_PORT", "5432", "数据库端口"],
        ["POSTGRES_USER", "ai_literacy", "数据库用户名"],
        ["POSTGRES_PASSWORD", "[建议改]", "数据库密码"],
        ["POSTGRES_DB", "ai_literacy_db", "数据库名"],
        ["ELASTICSEARCH_HOST", "elasticsearch", "搜索引擎地址"],
        ["ELASTICSEARCH_PORT", "9200", "搜索引擎端口"],
        ["MILVUS_HOST", "milvus-standalone", "向量数据库地址"],
        ["MILVUS_PORT", "19530", "向量数据库端口"],
        ["MINIO_HOST", "minio", "对象存储地址"],
        ["MINIO_PORT", "9000", "对象存储端口"],
        ["MINIO_ACCESS_KEY", "minioadmin", "对象存储访问密钥 [建议改]"],
        ["MINIO_SECRET_KEY", "minioadmin", "对象存储密码 [建议改]"],
        ["MINIO_BUCKET", "ai-literacy", "存储桶名称"],
        ["RABBITMQ_HOST", "rabbitmq", "消息队列地址"],
        ["RABBITMQ_PORT", "5672", "消息队列端口"],
        ["RABBITMQ_USER", "guest", "消息队列用户名"],
        ["RABBITMQ_PASSWORD", "guest", "消息队列密码 [建议改]"],
        ["REDIS_HOST", "redis", "缓存地址"],
        ["REDIS_PORT", "6379", "缓存端口"],
        ["LLM_API_KEY", "token-not-needed", "大模型API密钥"],
        ["LLM_BASE_URL", "http://host.docker.internal:8100/v1", "大模型API地址"],
        ["LLM_MODEL", "Qwen/Qwen3.5-35B-A3B", "大模型标识（与 --served-model-name 一致）"],
        ["JWT_SECRET_KEY", "[必改]", "JWT 密钥"],
        ["JWT_ALGORITHM", "HS256", "JWT 算法"],
        ["JWT_ACCESS_TOKEN_EXPIRE_MINUTES", "30", "Token 过期时间（分钟）"],
    ],
    col_widths=[5.5, 5.5, 5],
)


# ═══════════════════════════════════════════════════════════════════
# 保存文件
# ═══════════════════════════════════════════════════════════════════
output_path = os.path.join(os.path.dirname(__file__), "AI素养评测平台_系统管理员手册.docx")
doc.save(output_path)
print(f"文档已生成：{output_path}")
